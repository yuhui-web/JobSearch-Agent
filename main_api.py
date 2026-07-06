#!/usr/bin/env python
"""
JobSearch API - FastAPI application exposing the JobSearch-Agent functionality.

This API allows external applications (such as a React webapp) to communicate
with the JobSearch-Agent system through HTTP endpoints and WebSocket connections.
"""

import os
import sys
import json
import asyncio
import hashlib
import re
import base64
import tempfile
import threading
from collections import Counter
from io import BytesIO
from datetime import datetime
from typing import Dict, List, Any, Optional
from pathlib import Path

import uvicorn
from fastapi import FastAPI, WebSocket, BackgroundTasks, HTTPException, Depends, status
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse, Response
from fastapi.staticfiles import StaticFiles
from fastapi.security import APIKeyHeader
from pydantic import BaseModel, Field

# Fix for Windows Playwright subprocess issue in async context
if sys.platform == "win32":
    # Set Windows ProactorEventLoop policy to fix subprocess NotImplementedError
    asyncio.set_event_loop_policy(asyncio.WindowsProactorEventLoopPolicy())

    # Additional Windows subprocess environment setup
    os.environ["PLAYWRIGHT_BROWSERS_PATH"] = os.environ.get(
        "PLAYWRIGHT_BROWSERS_PATH", ""
    )

    # Ensure proper subprocess handling on Windows
    import signal

    if hasattr(signal, "SIGBREAK"):
        signal.signal(signal.SIGBREAK, signal.SIG_DFL)
    if hasattr(signal, "SIGINT"):
        signal.signal(signal.SIGINT, signal.SIG_DFL)

# Import agent functionality
# NOTE: Agent imports commented out for search-only testing
# from src.agents.job_details_parser import call_job_parsr_agent
# from src.agents.cv_writer import call_cv_agent
# from src.agents.coverLetter_writer import call_cover_letter_agent
from src.utils.job_search_pipeline import run_job_search, run_job_search_async
from src.utils.file_utils import slugify, ensure_dir_exists
from src.utils.job_database import JobDatabase
from src.utils.document_database import DocumentStorage, DocumentDatabase
from src.scraper.search.boss_scraper import (
    attach_company_map_url,
    build_smart_search_jobs,
    is_boss_automation_enabled,
    run_boss_deepseek_search,
    search_boss_jobs,
    search_observed_real_company_jobs,
    search_web_real_company_jobs,
)
from src.utils.deepseek_client import DeepSeekClient

# Create FastAPI app
app = FastAPI(
    title="JobSearch API",
    description="API for interacting with JobSearch-Agent system",
    version="1.0.0",
)

DOTENV_PATH = str(Path(__file__).resolve().parent / ".env")

try:
    from dotenv import load_dotenv

    load_dotenv(DOTENV_PATH, verbose=True, override=True)
except:
    pass

# API Key Authentication
DEFAULT_DEV_API_KEY = "dev-local-only-change-me"
ENVIRONMENT = os.getenv("ENVIRONMENT", "development").strip().lower()
API_KEY = os.getenv("API_KEY", DEFAULT_DEV_API_KEY)

if ENVIRONMENT in {"prod", "production"} and API_KEY == DEFAULT_DEV_API_KEY:
    raise RuntimeError("API_KEY must be set when ENVIRONMENT=production")

api_key_header = APIKeyHeader(name="X-API-Key", auto_error=False)


async def verify_api_key(api_key: str = Depends(api_key_header)):
    """Verify API key from request header"""
    if api_key != API_KEY:
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail="Invalid or missing API Key",
        )
    return api_key


def _parse_allowed_origins() -> List[str]:
    raw_origins = os.getenv(
        "ALLOWED_ORIGIN",
        "http://127.0.0.1:5173,http://localhost:5173,http://127.0.0.1:5174,http://localhost:5174",
    )
    origins = [origin.strip() for origin in raw_origins.split(",") if origin.strip()]
    if ENVIRONMENT in {"prod", "production"} and "*" in origins:
        raise RuntimeError("Wildcard CORS origin is not allowed in production")
    return origins or ["http://127.0.0.1:5173"]


# Add CORS middleware for React frontend
app.add_middleware(
    CORSMiddleware,
    allow_origins=_parse_allowed_origins(),
    allow_credentials=False,
    allow_methods=["GET", "POST", "DELETE", "OPTIONS"],
    allow_headers=["Content-Type", "X-API-Key"],
)

# Make output directory available for downloading generated files
output_dir = os.path.join(os.getcwd(), "output")
ensure_dir_exists(output_dir)
app.mount("/output", StaticFiles(directory=output_dir), name="output")


# Define data models
class JobSearchRequest(BaseModel):
    keywords: str
    locations: List[str] = Field(default=["Remote"])
    job_type: str = Field(default="full-time")
    experience_level: str = Field(default="mid-level")
    max_jobs: int = Field(default=3)
    scrapers: List[str] = Field(
        default=["linkedin"]
    )  # Scraper selection (only LinkedIn working currently)
    candidate_profile: Optional[str] = None


class BossMonitorRequest(JobSearchRequest):
    interval_seconds: int = Field(default=300, ge=60)


class ImportedJobsRequest(BaseModel):
    jobs: List[Dict[str, Any]]
    keywords: str = Field(default="")
    locations: List[str] = Field(default_factory=list)
    job_type: str = Field(default="internship")
    experience_level: str = Field(default="entry-level")
    max_jobs: int = Field(default=20)
    candidate_profile: Optional[str] = None
    source: str = Field(default="boss_import")


class BossTextAnalyzeRequest(BaseModel):
    job_text: str
    source_url: Optional[str] = None
    candidate_profile: Optional[str] = None


class CareerAnalyzeRequest(BaseModel):
    candidate_profile: str
    target_role: str = Field(default="Python实习")
    location: str = Field(default="武汉")
    job_type: str = Field(default="internship")
    experience_level: str = Field(default="entry-level")
    requirement_text: Optional[str] = None
    max_recommendations: int = Field(default=5)


class ResumeExtractRequest(BaseModel):
    file_name: str
    content_base64: str


class JobParseRequest(BaseModel):
    text: Optional[str] = None
    file_content: Optional[str] = None
    url: Optional[str] = None
    extract_webpage: bool = Field(default=False)


class JobProcessRequest(BaseModel):
    job_posting: Dict[str, Any]
    generate_cv: bool = Field(default=True)
    generate_cover_letter: bool = Field(default=False)


class InterviewLogRequest(BaseModel):
    job_id: Optional[int] = None
    job_title: str
    company_name: str
    interview_date: Optional[str] = None
    outcome: str
    failure_reason: Optional[str] = None
    notes: Optional[str] = None
    next_action: Optional[str] = None


class WebSocketMessage(BaseModel):
    action: str
    data: Dict[str, Any]


# Store active websocket connections
active_connections: List[WebSocket] = []


# WebSocket connection manager
class ConnectionManager:
    def __init__(self):
        self.active_connections: List[WebSocket] = []

    async def connect(self, websocket: WebSocket):
        await websocket.accept()
        self.active_connections.append(websocket)

    def disconnect(self, websocket: WebSocket):
        if websocket in self.active_connections:
            self.active_connections.remove(websocket)

    async def send_progress(self, websocket: WebSocket, message: str):
        if websocket in self.active_connections:
            await websocket.send_json({"type": "progress", "message": message})

    async def send_result(self, websocket: WebSocket, data: Any):
        if websocket in self.active_connections:
            await websocket.send_json({"type": "result", "data": data})

    async def send_error(self, websocket: WebSocket, error: str):
        if websocket in self.active_connections:
            await websocket.send_json({"type": "error", "message": error})

    async def broadcast(self, payload: Dict[str, Any]):
        disconnected: List[WebSocket] = []
        for websocket in list(self.active_connections):
            try:
                await websocket.send_json(payload)
            except Exception:
                disconnected.append(websocket)
        for websocket in disconnected:
            self.disconnect(websocket)


manager = ConnectionManager()

boss_monitor_task: Optional[asyncio.Task] = None
boss_monitor_state: Dict[str, Any] = {
    "running": False,
    "keywords": "",
    "locations": [],
    "job_type": "",
    "experience_level": "",
    "max_jobs": 0,
    "scrapers": ["boss"],
    "interval_seconds": 300,
    "last_search_id": None,
    "last_job_count": 0,
    "last_run_at": None,
    "last_error": None,
}


# Search history tracking
search_history_file = os.path.join(output_dir, "search_history.json")
search_history_lock = threading.RLock()


def load_search_history() -> List[Dict[str, Any]]:
    """Load search history from file"""
    with search_history_lock:
        if os.path.exists(search_history_file):
            try:
                with open(search_history_file, "r", encoding="utf-8") as f:
                    return json.load(f)
            except Exception as e:
                print(f"Error loading search history: {e}")
                return []
    return []


def save_search_history(history: List[Dict[str, Any]]):
    """Save search history to file"""
    with search_history_lock:
        tmp_path = None
        try:
            ensure_dir_exists(output_dir)
            fd, tmp_path = tempfile.mkstemp(
                prefix="search_history_",
                suffix=".json",
                dir=output_dir,
            )
            with os.fdopen(fd, "w", encoding="utf-8") as f:
                json.dump(history, f, indent=2, ensure_ascii=False)
            os.replace(tmp_path, search_history_file)
        except Exception as e:
            print(f"Error saving search history: {e}")
            if tmp_path and os.path.exists(tmp_path):
                try:
                    os.remove(tmp_path)
                except OSError:
                    pass


def generate_search_hash(
    keywords: str,
    locations: List[str],
    job_type: str,
    experience_level: str,
    scrapers: List[str],
) -> str:
    """Generate a hash for search parameters to identify duplicates"""
    # Normalize parameters for consistent hashing
    normalized_locations = sorted([loc.lower().strip() for loc in locations])
    normalized_scrapers = sorted([s.lower().strip() for s in scrapers])

    search_string = f"{keywords.lower().strip()}|{','.join(normalized_locations)}|{job_type.lower().strip()}|{experience_level.lower().strip()}|{','.join(normalized_scrapers)}"
    return hashlib.md5(search_string.encode()).hexdigest()


def find_similar_searches(
    keywords: str,
    locations: List[str],
    job_type: str,
    experience_level: str,
    scrapers: List[str],
) -> Dict[str, List[Dict[str, Any]]]:
    """Find similar searches in history"""
    search_hash = generate_search_hash(
        keywords, locations, job_type, experience_level, scrapers
    )
    history = load_search_history()

    # Find exact matches
    exact_matches = [
        search for search in history if search.get("search_hash") == search_hash
    ]

    # Find partial matches (same keywords and job type)
    partial_matches = [
        search
        for search in history
        if search.get("keywords", "").lower().strip() == keywords.lower().strip()
        and search.get("job_type", "").lower().strip() == job_type.lower().strip()
        and search.get("search_hash") != search_hash
    ]

    return {"exact": exact_matches, "similar": partial_matches}


def add_search_to_history(
    search_id: str,
    keywords: str,
    locations: List[str],
    job_type: str,
    experience_level: str,
    scrapers: List[str],
    max_jobs: int,
    status: str = "started",
):
    """Add a search to history"""
    with search_history_lock:
        history = load_search_history()
        search_hash = generate_search_hash(
            keywords, locations, job_type, experience_level, scrapers
        )

        search_entry = {
            "search_id": search_id,
            "search_hash": search_hash,
            "keywords": keywords,
            "locations": locations,
            "job_type": job_type,
            "experience_level": experience_level,
            "scrapers": scrapers,
            "max_jobs": max_jobs,
            "timestamp": datetime.now().isoformat(),
            "status": status,
        }

        # Add to beginning of history (most recent first)
        history.insert(0, search_entry)

        # Keep only last 50 searches
        history = history[:50]

        save_search_history(history)


def update_search_status(search_id: str, status: str, job_count: int = None):
    """Update search status in history"""
    with search_history_lock:
        history = load_search_history()
        for search in history:
            if search.get("search_id") == search_id:
                search["status"] = status
                if job_count is not None:
                    search["job_count"] = job_count
                break
        save_search_history(history)


def boss_monitor_snapshot() -> Dict[str, Any]:
    """Return a serializable copy of the current BOSS monitor state."""
    snapshot = dict(boss_monitor_state)
    history = load_search_history()
    latest_monitor = next(
        (
            search
            for search in history
            if str(search.get("search_id", "")).startswith("boss_monitor_")
        ),
        None,
    )
    if latest_monitor and (
        not snapshot.get("last_search_id") or int(snapshot.get("last_job_count") or 0) == 0
    ):
        snapshot.update(
            {
                "keywords": latest_monitor.get("keywords", snapshot.get("keywords", "")),
                "locations": latest_monitor.get("locations", snapshot.get("locations", [])),
                "job_type": latest_monitor.get("job_type", snapshot.get("job_type", "")),
                "experience_level": latest_monitor.get(
                    "experience_level", snapshot.get("experience_level", "")
                ),
                "max_jobs": latest_monitor.get("max_jobs", snapshot.get("max_jobs", 0)),
                "scrapers": latest_monitor.get("scrapers", snapshot.get("scrapers", ["boss"])),
                "last_search_id": latest_monitor.get("search_id"),
                "last_job_count": latest_monitor.get("job_count", 0),
                "last_run_at": latest_monitor.get("timestamp", snapshot.get("last_run_at")),
            }
        )
    return snapshot


def run_boss_monitor_tick(config: Dict[str, Any]) -> Dict[str, Any]:
    """Fetch one BOSS monitor batch and persist it as a normal search result."""
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S_%f")
    search_id = f"boss_monitor_{timestamp}"
    keywords = str(config.get("keywords") or "").strip()
    locations = list(config.get("locations") or [])
    job_type = str(config.get("job_type") or "internship")
    experience_level = str(config.get("experience_level") or "entry-level")
    max_jobs = int(config.get("max_jobs") or 5)
    scrapers = list(config.get("scrapers") or ["boss"])
    candidate_profile = config.get("candidate_profile")

    boss_monitor_state["last_run_at"] = datetime.now().isoformat()
    boss_monitor_state["last_error"] = None

    try:
        results = run_boss_deepseek_search(
            keywords=keywords,
            locations=locations,
            max_jobs=max_jobs,
            candidate_profile=candidate_profile,
            job_type=job_type,
            experience_level=experience_level,
        )
        results = [
            {
                **job,
                "monitor_source": job.get("source"),
                "source": "boss_monitor",
            }
            for job in results
        ]
        empty_note = None
        if not results and not is_boss_automation_enabled():
            empty_note = (
                "BOSS automation is not enabled; monitor could not pull jobs from the logged-in BOSS page."
            )
        result_file = os.path.join(output_dir, f"{search_id}.json")
        with open(result_file, "w", encoding="utf-8") as f:
            json.dump(results, f, indent=2, ensure_ascii=False)

        add_search_to_history(
            search_id=search_id,
            keywords=keywords,
            locations=locations,
            job_type=job_type,
            experience_level=experience_level,
            scrapers=scrapers,
            max_jobs=max_jobs,
            status="completed",
        )
        update_search_status(search_id, status="completed", job_count=len(results))
        boss_monitor_state.update(
            {
                "last_search_id": search_id,
                "last_job_count": len(results),
                "last_error": empty_note,
            }
        )
        return {
            "type": "boss_monitor_update",
            "search_id": search_id,
            "status": "completed",
            "job_count": len(results),
            "updated_at": boss_monitor_state["last_run_at"],
            "message": empty_note,
        }
    except Exception as exc:
        boss_monitor_state["last_error"] = str(exc)
        error_file = os.path.join(output_dir, f"{search_id}_error.txt")
        with open(error_file, "w", encoding="utf-8") as f:
            f.write(str(exc))
        return {
            "type": "boss_monitor_update",
            "search_id": search_id,
            "status": "error",
            "job_count": 0,
            "updated_at": boss_monitor_state["last_run_at"],
            "error": str(exc),
        }


async def boss_monitor_loop(config: Dict[str, Any]):
    """Run BOSS monitor ticks until stopped."""
    interval_seconds = max(60, int(config.get("interval_seconds") or 300))
    loop = asyncio.get_running_loop()
    try:
        while boss_monitor_state.get("running"):
            result = await loop.run_in_executor(None, run_boss_monitor_tick, config)
            await manager.broadcast({"type": "boss_monitor_update", "data": result})
            await asyncio.sleep(interval_seconds)
    except asyncio.CancelledError:
        raise
    finally:
        boss_monitor_state["running"] = False


def delete_search_artifacts(search_id: str) -> int:
    """Delete output artifacts for a search id and return removed file count."""
    removed = 0
    for suffix in (".json", "_error.txt"):
        path = os.path.join(output_dir, f"{search_id}{suffix}")
        if os.path.exists(path):
            os.remove(path)
            removed += 1
    return removed


def get_job_database() -> JobDatabase:
    """Create a job database connection for endpoint handlers."""
    return JobDatabase()


def _normalize_search_text(value: Any) -> str:
    """Normalize user/search text so Chinese and English spacing differences still match."""
    return re.sub(r"\s+", "", str(value or "").lower())


def _job_title(job: Dict[str, Any]) -> Any:
    return job.get("job_title") or job.get("title") or job.get("position_title")


def _company_name(job: Dict[str, Any]) -> Any:
    return job.get("company_name") or job.get("company") or job.get("company_title")


def _job_location(job: Dict[str, Any]) -> Any:
    return job.get("job_location") or job.get("location") or job.get("work_location")


def _job_matches_keyword(job: Dict[str, Any], keywords: str) -> bool:
    needle = _normalize_search_text(keywords)
    if not needle:
        return True
    haystack = _normalize_search_text(
        " ".join(
            str(value or "")
            for value in [
                _job_title(job),
                job.get("job_description"),
                _company_name(job),
                " ".join(str(tag) for tag in job.get("tags", []) or []),
            ]
        )
    )
    if needle in haystack:
        return True
    tokens = re.findall(r"[a-z0-9]+|[\u4e00-\u9fff]+", str(keywords or "").lower())
    tokens = [token for token in tokens if len(token) > 1 or re.fullmatch(r"[a-z0-9]+", token)]
    if job.get("source") == "boss_import":
        generic_tokens = {
            "实习", "实习生", "初级", "入门", "开发", "工程师", "后端", "前端", "岗位", "校招",
            "intern", "internship", "entry", "level", "junior",
        }
        role_tokens = [token for token in tokens if _normalize_search_text(token) not in generic_tokens]
        if role_tokens:
            return any(_normalize_search_text(token) in haystack for token in role_tokens)
        return bool(tokens) and any(_normalize_search_text(token) in haystack for token in tokens)
    return bool(tokens) and all(token in haystack for token in tokens)


def _location_parts(value: Any) -> List[str]:
    text = str(value or "").strip()
    if not text:
        return []
    return [
        part.strip()
        for part in re.split(r"[\s,，、/|;；·\-.]+", text)
        if part.strip()
    ]


def _extract_requested_districts(locations: List[str]) -> List[str]:
    districts: List[str] = []
    suffixes = (
        "\u533a",
        "\u53bf",
        "\u65d7",
        "\u65b0\u533a",
        "\u5f00\u53d1\u533a",
    )
    for location in locations or []:
        for part in _location_parts(location):
            if any(suffix in part for suffix in suffixes):
                districts.append(part)
    return districts


def _job_matches_specific_district(job: Dict[str, Any], locations: List[str]) -> bool:
    requested_districts = [_normalize_search_text(district) for district in _extract_requested_districts(locations)]
    if not requested_districts:
        return _job_matches_locations(job, locations)
    job_location = _normalize_search_text(_job_location(job))
    return any(district in job_location for district in requested_districts)


def _job_matches_locations(job: Dict[str, Any], locations: List[str]) -> bool:
    requested_locations = [_normalize_search_text(location) for location in locations if location]
    if not requested_locations:
        return True
    job_location = _normalize_search_text(_job_location(job))
    return any(location in job_location or job_location in location for location in requested_locations)


def _has_specific_district(locations: List[str]) -> bool:
    return bool(_extract_requested_districts(locations))


def _filter_jobs_for_specific_location(jobs: List[Dict[str, Any]], locations: List[str]) -> List[Dict[str, Any]]:
    if not _has_specific_district(locations):
        return jobs
    return [job for job in jobs if _job_matches_specific_district(job, locations)]


def _enrich_fast_candidates_with_analysis(
    jobs: List[Dict[str, Any]], candidate_profile: Optional[str]
) -> List[Dict[str, Any]]:
    analyzer = DeepSeekClient()
    enriched: List[Dict[str, Any]] = []
    for job in jobs:
        hydrated = dict(job)
        if "ai_analysis" not in hydrated:
            hydrated["ai_analysis"] = analyzer.analyze_job(hydrated, candidate_profile=candidate_profile)
        enriched.append(hydrated)

    return sorted(
        enriched,
        key=lambda job: job.get("ai_analysis", {}).get("match_score") or 0,
        reverse=True,
    )


def _filter_jobs_for_search(
    jobs: List[Dict[str, Any]],
    keywords: str,
    locations: List[str],
    max_jobs: Optional[int] = None,
) -> List[Dict[str, Any]]:
    blocked = []
    exact_location_matches = []
    relaxed_location_matches = []
    for job in jobs:
        if job.get("ai_analysis", {}).get("status") == "blocked":
            blocked.append(job)
            continue
        if job.get("source") in {"smart_search", "curated"}:
            continue
        if not JobDatabase.has_valid_job_identity(_job_title(job), _company_name(job)):
            continue
        if not _job_matches_keyword(job, keywords):
            continue
        if _job_matches_locations(job, locations):
            exact_location_matches.append(job)
        else:
            # District selection is a priority, not a hard empty-state trap:
            # when the chosen district has too few candidates, keep same-role
            # company candidates so a requested list of 5 can still render.
            relaxed_location_matches.append(job)

    filtered = blocked + exact_location_matches + relaxed_location_matches
    if max_jobs is not None:
        return filtered[:max_jobs]
    return filtered


def _trend_bucket(counter: Counter, limit: int = 8) -> List[Dict[str, Any]]:
    return [
        {"name": name, "count": count}
        for name, count in counter.most_common(limit)
        if str(name).strip()
    ]


def _detect_job_keywords(job: Dict[str, Any]) -> List[str]:
    text = _normalize_search_text(
        " ".join(
            str(value or "")
            for value in [
                _job_title(job),
                job.get("job_description"),
                job.get("tags"),
                job.get("salary"),
            ]
        )
    )
    patterns = [
        ("Python", r"python|fastapi|flask|pandas|爬虫|数据处理|数据分析"),
        ("Java", r"java|springboot|spring|mybatis"),
        ("前端", r"vue|react|javascript|typescript|html|css|前端"),
        ("AI Agent", r"agent|llm|rag|大模型|智能体"),
        ("数据库", r"mysql|oracle|redis|sql|数据库"),
        ("C/C++", r"c\+\+|cpp|c语言|嵌入式"),
        ("测试/自动化", r"测试|自动化|selenium|playwright|接口"),
    ]
    return [label for label, pattern in patterns if re.search(pattern, text)]


def _detect_company_type(company_name: Any, job: Dict[str, Any]) -> str:
    text = str(company_name or "") + " " + str(job.get("job_description") or "")
    if re.search(r"外包|人力|服务外包|派遣", text):
        return "外包/人力"
    if re.search(r"汽车|车|智能驾驶|新能源", text):
        return "汽车/智能硬件"
    if re.search(r"教育|培训|学校|学院", text):
        return "教育培训"
    if re.search(r"软件|信息|科技|智能|网络|数据|云", text):
        return "软件服务"
    return "普通企业"


def build_job_trends(
    jobs: List[Dict[str, Any]],
    target_role: str = "",
    location: str = "",
) -> Dict[str, Any]:
    """Aggregate visible job candidates into a market trend snapshot."""
    keyword_counter: Counter = Counter()
    city_counter: Counter = Counter()
    salary_counter: Counter = Counter()
    company_type_counter: Counter = Counter()

    for job in jobs or []:
        for keyword in _detect_job_keywords(job):
            keyword_counter[keyword] += 1
        if _job_location(job):
            city_counter[str(_job_location(job)).strip()] += 1
        if job.get("salary"):
            salary_counter[str(job.get("salary")).strip()] += 1
        company_type_counter[_detect_company_type(_company_name(job), job)] += 1

    top_keywords = _trend_bucket(keyword_counter)
    top_locations = _trend_bucket(city_counter)
    top_salaries = _trend_bucket(salary_counter)
    top_company_types = _trend_bucket(company_type_counter)

    if jobs:
        leading = top_keywords[0]["name"] if top_keywords else (target_role or "目标岗位")
        place = top_locations[0]["name"] if top_locations else (location or "目标城市")
        salary = top_salaries[0]["name"] if top_salaries else "薪资待核验"
        summary = f"市场最近更关注 {leading}，主要集中在 {place}，常见薪资为 {salary}。"
    else:
        summary = "还没有真实岗位样本，无法形成可靠市场趋势。"

    return {
        "total_jobs": len(jobs or []),
        "summary": summary,
        "keywords": top_keywords,
        "cities": top_locations,
        "salary_ranges": top_salaries,
        "company_types": top_company_types,
    }


def _extract_labeled_value(text: str, labels: List[str]) -> Optional[str]:
    for label in labels:
        pattern = rf"{re.escape(label)}\s*[:：]\s*(.+)"
        match = re.search(pattern, text, flags=re.I)
        if match:
            return match.group(1).strip()
    return None


def parse_boss_job_text(job_text: str, source_url: Optional[str] = None) -> Dict[str, Any]:
    """Turn pasted BOSS/JD text into the job dict shape used by the Agent flow."""
    normalized_lines = [line.strip() for line in job_text.splitlines() if line.strip()]
    joined_text = "\n".join(normalized_lines).strip()
    first_line = normalized_lines[0] if normalized_lines else "未命名岗位"

    job_title = _extract_labeled_value(joined_text, ["职位", "岗位", "岗位名称", "职位名称"])
    company_name = _extract_labeled_value(joined_text, ["公司", "公司名称", "企业", "企业名称"])
    job_location = _extract_labeled_value(joined_text, ["地点", "位置", "工作地点", "城市"])
    salary = _extract_labeled_value(joined_text, ["薪资", "工资", "待遇"])

    if not job_title:
        job_title = first_line[:80]
    if not company_name:
        company_name = "待识别公司"
    if not job_location:
        location_match = re.search(r"(武汉|北京|上海|广州|深圳|杭州|成都|远程)[^\n，,；; ]*", joined_text)
        job_location = location_match.group(0) if location_match else ""
    if not salary:
        salary_match = re.search(r"\d+\s*[-~至]\s*\d+\s*(?:元/天|/天|K|k|千|万|元)", joined_text)
        salary = salary_match.group(0) if salary_match else ""

    return {
        "source": "boss_text",
        "source_url": source_url,
        "scraped_at": datetime.now().isoformat(timespec="seconds"),
        "job_title": job_title,
        "company_name": company_name,
        "job_location": job_location,
        "salary": salary,
        "job_description": joined_text,
    }


def extract_resume_text(file_name: str, content: bytes) -> str:
    suffix = Path(file_name).suffix.lower()
    if suffix == ".docx":
        try:
            from docx import Document
        except Exception as exc:
            raise ValueError("python-docx 未安装，无法解析 Word 简历。") from exc

        document = Document(BytesIO(content))
        paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
        table_cells = [
            cell.text.strip()
            for table in document.tables
            for row in table.rows
            for cell in row.cells
            if cell.text.strip()
        ]
        return "\n".join(paragraphs + table_cells).strip()

    if suffix in {".txt", ".md", ".markdown", ".json"}:
        for encoding in ("utf-8-sig", "utf-8", "gb18030"):
            try:
                return content.decode(encoding).strip()
            except UnicodeDecodeError:
                continue
        raise ValueError("无法识别文本简历编码，请另存为 UTF-8 后重试。")

    raise ValueError("目前支持 .docx、.txt、.md、.markdown、.json 简历文件。")


def _find_search_history_entry(search_id: str) -> Optional[Dict[str, Any]]:
    for search in load_search_history():
        if search.get("search_id") == search_id:
            return search
    return None


def _hydrate_job_json_fields(job: Dict[str, Any]) -> Dict[str, Any]:
    hydrated = dict(job)
    for field in [
        "job_insights",
        "apply_info",
        "company_info",
        "hiring_team",
        "related_jobs",
    ]:
        if hydrated.get(field):
            try:
                hydrated[field] = json.loads(hydrated[field])
            except (json.JSONDecodeError, TypeError):
                pass
    return hydrated


def _first_non_empty(job: Dict[str, Any], *fields: str) -> str:
    for field in fields:
        value = job.get(field)
        if value is not None and str(value).strip():
            return str(value).strip()
    return ""


def sanitize_salary(value: Any) -> str:
    """Hide BOSS icon-font salary text that cannot be rendered as real numbers."""
    text = str(value or "").strip()
    if not text:
        return ""
    has_private_use_chars = any("\ue000" <= char <= "\uf8ff" for char in text)
    has_tofu_chars = "□" in text or "\ufffd" in text
    has_unknown_runs = text.count("?") >= 2
    if has_private_use_chars or has_tofu_chars or has_unknown_runs:
        return "薪资到 BOSS 核验"
    return text


def sanitize_job_display_fields(job: Dict[str, Any]) -> Dict[str, Any]:
    cleaned = dict(job)
    cleaned["salary"] = sanitize_salary(cleaned.get("salary"))
    title = _job_title(cleaned)
    company = _company_name(cleaned)
    location = _job_location(cleaned)
    if title:
        cleaned["job_title"] = title
        cleaned["title"] = title
        cleaned["name"] = title
    if company:
        cleaned["company_name"] = company
        cleaned["company"] = company
    if location:
        cleaned["job_location"] = location
        cleaned["location"] = location
    return cleaned


def normalize_imported_job(job: Dict[str, Any], default_location: str = "") -> Dict[str, Any]:
    """Normalize a BOSS page card collected in-browser into the app's job shape."""
    title = _first_non_empty(job, "job_title", "title", "name", "position", "positionName")
    company = _first_non_empty(job, "company_name", "company", "companyName", "brandName")
    location = _first_non_empty(job, "job_location", "location", "area", "city", "workAddress") or default_location
    salary = _first_non_empty(job, "salary", "pay", "wage")
    source_url = _first_non_empty(job, "source_url", "url", "link", "href", "detailUrl")
    tags = job.get("tags") if isinstance(job.get("tags"), list) else []
    description = _first_non_empty(job, "job_description", "description", "jd", "detail", "text")
    if not description:
        description = " ".join(str(item) for item in [title, company, location, salary, " ".join(tags)] if item)

    normalized = {
        "source": "boss_import",
        "source_url": source_url,
        "link_status": "verified_detail" if "/job_detail/" in source_url else "needs_verification",
        "scraped_at": datetime.now().isoformat(timespec="seconds"),
        "job_title": title,
        "company_name": company,
        "job_location": location,
        "salary": sanitize_salary(salary),
        "job_description": description,
        "tags": tags,
    }
    if job.get("id"):
        normalized["external_id"] = str(job["id"])
    if job.get("workAddress") or job.get("address"):
        normalized["work_address"] = _first_non_empty(job, "workAddress", "address")
    return sanitize_job_display_fields(repair_imported_job_identity(normalized))


def _looks_like_location_name(value: Any) -> bool:
    text = str(value or "").strip()
    if not text:
        return False
    company_markers = (
        "科技", "信息", "智能", "软件", "网络", "数据", "教育", "传媒", "汽车",
        "有限公司", "集团", "工作室", "中心", "咨询", "服务", "电子", "医疗",
    )
    if any(marker in text for marker in company_markers):
        return False
    location_markers = (
        "武汉", "洪山", "江夏", "江汉", "江岸", "汉阳", "武昌", "硚口",
        "东西湖", "青山", "东湖", "蔡甸", "黄陂", "新洲", "经开", "光谷",
        "南昌", "东湖区", "江岸区", "洪山区", "区", "县", "市", "路", "街", "常青",
    )
    return "·" in text or any(marker in text for marker in location_markers)


def _looks_like_job_title(value: Any) -> bool:
    text = str(value or "").strip().lower()
    if not text:
        return False
    title_markers = (
        "实习", "工程师", "开发", "测试", "算法", "研究员", "助理", "全栈",
        "后端", "前端", "python", "java", "agent", "llm", "rag", "fastapi",
    )
    return any(marker in text for marker in title_markers)


def _pick_company_from_imported_tags(job: Dict[str, Any]) -> str:
    tags = job.get("tags") if isinstance(job.get("tags"), list) else []
    for item in tags:
        candidate = str(item or "").strip()
        if not candidate:
            continue
        if _looks_like_location_name(candidate) or _looks_like_job_title(candidate):
            continue
        if len(candidate) > 2:
            return candidate
    return ""


def repair_imported_job_identity(job: Dict[str, Any]) -> Dict[str, Any]:
    """Fix common BOSS collector mistakes, especially company/location swaps."""
    repaired = dict(job)
    company = _company_name(repaired)
    if not company or _looks_like_location_name(company):
        tag_company = _pick_company_from_imported_tags(repaired)
        if tag_company:
            repaired["company_name"] = tag_company

    location = _job_location(repaired)
    if not location and company and _looks_like_location_name(company):
        repaired["job_location"] = company
    return repaired


def _looks_corrupted_text(value: Any) -> bool:
    text = str(value or "").strip()
    return text.count("?") >= 2 or "\ufffd" in text


def save_imported_jobs_as_search(request: ImportedJobsRequest) -> Dict[str, Any]:
    default_location = (request.locations or [""])[0]
    is_monitor_import = str(request.source or "").startswith("boss_monitor")
    selected: List[Dict[str, Any]] = []
    seen: set[tuple[str, str]] = set()

    for raw_job in request.jobs:
        normalized = normalize_imported_job(raw_job, default_location=default_location)
        identity = (
            _normalize_search_text(_job_title(normalized)),
            _normalize_search_text(_company_name(normalized)),
        )
        if not all(identity) or identity in seen:
            continue
        if not JobDatabase.has_valid_job_identity(_job_title(normalized), _company_name(normalized)):
            continue
        if is_monitor_import:
            normalized["monitor_source"] = normalized.get("source")
            normalized["source"] = "boss_monitor"
        seen.add(identity)
        selected.append(attach_company_map_url(normalized, default_location))
        if len(selected) >= max(1, request.max_jobs):
            break

    enriched = _enrich_fast_candidates_with_analysis(selected, request.candidate_profile)
    search_prefix = "boss_monitor" if is_monitor_import else "job_import"
    search_id = f"{search_prefix}_{datetime.now().strftime('%Y%m%d_%H%M%S_%f')}"
    result_file = os.path.join(output_dir, f"{search_id}.json")
    with open(result_file, "w", encoding="utf-8") as f:
        json.dump(enriched, f, indent=2, ensure_ascii=False)

    add_search_to_history(
        search_id=search_id,
        keywords=request.keywords or "Imported BOSS jobs",
        locations=request.locations or ([default_location] if default_location else []),
        job_type=request.job_type,
        experience_level=request.experience_level,
        scrapers=[request.source or "boss_import"],
        max_jobs=request.max_jobs,
        status="completed",
    )
    update_search_status(search_id, status="completed", job_count=len(enriched))
    if is_monitor_import:
        boss_monitor_state.update(
            {
                "running": True,
                "keywords": request.keywords or boss_monitor_state.get("keywords", ""),
                "locations": request.locations or boss_monitor_state.get("locations", []),
                "job_type": request.job_type,
                "experience_level": request.experience_level,
                "max_jobs": request.max_jobs,
                "scrapers": [request.source],
                "last_search_id": search_id,
                "last_job_count": len(enriched),
                "last_run_at": datetime.now().isoformat(),
                "last_error": None,
            }
        )
    return {"search_id": search_id, "status": "Imported jobs saved", "job_count": len(enriched)}


def load_recent_imported_jobs(
    keywords: str,
    locations: List[str],
    max_jobs: int,
    candidate_profile: Optional[str] = None,
    strict_location: bool = False,
) -> List[Dict[str, Any]]:
    """Use recently imported real jobs as the first-class real source."""
    max_age_seconds = int(os.getenv("RECENT_IMPORTED_JOB_MAX_AGE_SECONDS", "3600"))
    now_ts = datetime.now().timestamp()
    imported_files = sorted(
        Path(output_dir).glob("job_import_*.json"),
        key=lambda path: path.stat().st_mtime,
        reverse=True,
    )
    selected: List[Dict[str, Any]] = []
    seen: set[tuple[str, str]] = set()

    for imported_file in imported_files:
        try:
            if max_age_seconds > 0 and now_ts - imported_file.stat().st_mtime > max_age_seconds:
                continue
        except OSError:
            continue
        try:
            jobs = json.loads(imported_file.read_text(encoding="utf-8"))
        except (OSError, json.JSONDecodeError):
            continue
        if not isinstance(jobs, list):
            continue

        for raw_job in jobs:
            if not isinstance(raw_job, dict):
                continue
            if raw_job.get("source") != "boss_import":
                continue
            job = sanitize_job_display_fields(repair_imported_job_identity(raw_job))
            identity = (
                _normalize_search_text(_job_title(job)),
                _normalize_search_text(_company_name(job)),
            )
            if not all(identity) or identity in seen:
                continue
            if _looks_corrupted_text(_job_title(job)) or _looks_corrupted_text(_company_name(job)):
                continue
            if not JobDatabase.has_valid_job_identity(_job_title(job), _company_name(job)):
                continue
            seen.add(identity)
            selected.append(job)

    if not selected:
        return []

    if strict_location and _has_specific_district(locations):
        selected = _filter_jobs_for_specific_location(selected, locations)

    filtered = _filter_jobs_for_search(selected, keywords, locations, max_jobs=max(max_jobs * 3, max_jobs))
    default_location = (locations or ["武汉"])[0]
    enriched = [
        attach_company_map_url(sanitize_job_display_fields(repair_imported_job_identity(job)), default_location)
        for job in filtered[:max_jobs]
    ]
    return _enrich_fast_candidates_with_analysis(enriched, candidate_profile)


def load_search_results_from_database(
    keywords: str,
    locations: List[str],
    max_jobs: int,
) -> List[Dict[str, Any]]:
    """Load only valid database jobs matching the current search criteria."""
    db = JobDatabase()
    try:
        candidate_jobs = db.get_jobs(limit=1000)
        results = []
        for job in candidate_jobs:
            filtered_jobs = _filter_jobs_for_search([job], keywords, locations)
            if not filtered_jobs:
                continue
            results.append(_hydrate_job_json_fields(filtered_jobs[0]))
            if len(results) >= max_jobs:
                break
        return results
    finally:
        db.close()


def regenerate_visible_search_results(search: Dict[str, Any]) -> List[Dict[str, Any]]:
    """Rebuild visible BOSS/Amap candidates when an old result file filters to zero."""
    jobs = build_fast_boss_candidates(
        keywords=search.get("keywords", ""),
        locations=search.get("locations", []) or ["武汉"],
        max_jobs=int(search.get("max_jobs") or 5),
        candidate_profile=search.get("candidate_profile"),
        job_type=search.get("job_type", "internship"),
        experience_level=search.get("experience_level", "entry-level"),
    )
    result_file = os.path.join(output_dir, f"{search['search_id']}.json")
    with open(result_file, "w", encoding="utf-8") as f:
        json.dump(jobs, f, indent=2, ensure_ascii=False)
    update_search_status(search["search_id"], status="completed", job_count=len(jobs))
    return jobs


def build_fast_boss_candidates(
    keywords: str,
    locations: List[str],
    max_jobs: int,
    candidate_profile: Optional[str] = None,
    job_type: str = "internship",
    experience_level: str = "entry-level",
    strict_location: bool = False,
) -> List[Dict[str, Any]]:
    """Return only jobs collected from live/real search sources.

    Local samples and generated company candidates are intentionally excluded
    from this path. If live sources are blocked or empty, the caller should show
    an honest empty/blocked state instead of filling the UI with fake results.
    """
    automation_enabled = is_boss_automation_enabled()
    live_jobs: List[Dict[str, Any]] = []
    if automation_enabled:
        try:
            live_jobs.extend(search_boss_jobs(keywords, locations, max_jobs))
        except Exception:
            # BOSS often redirects to login/security verification. In strict mode
            # that means "no BOSS result", not "invent a fallback result".
            pass

    if len(live_jobs) < max_jobs:
        live_jobs.extend(
            search_web_real_company_jobs(
                keywords=keywords,
                locations=locations,
                max_jobs=max_jobs - len(live_jobs),
                candidate_profile=candidate_profile,
                job_type=job_type,
                experience_level=experience_level,
                search_plan=[],
            )
        )

    jobs = _filter_jobs_for_specific_location(live_jobs, locations)
    jobs = [
        job for job in jobs
        if job.get("source") in {"boss", "web_search"}
        and JobDatabase.has_valid_job_identity(_job_title(job), _company_name(job))
    ]
    if jobs:
        jobs = [attach_company_map_url(job, (locations or ["武汉"])[0]) for job in jobs[:max_jobs]]
        return _enrich_fast_candidates_with_analysis(jobs, candidate_profile)

    imported_jobs = load_recent_imported_jobs(
        keywords=keywords,
        locations=locations,
        max_jobs=max_jobs,
        candidate_profile=candidate_profile,
        strict_location=True,
    )
    return imported_jobs[:max_jobs]


def build_legacy_candidate_jobs(
    keywords: str,
    locations: List[str],
    max_jobs: int,
    candidate_profile: Optional[str] = None,
    job_type: str = "internship",
    experience_level: str = "entry-level",
) -> List[Dict[str, Any]]:
    """Legacy non-live candidate generation kept out of the strict search path."""
    def normalize_candidate_source(job: Dict[str, Any]) -> Dict[str, Any]:
        hydrated = dict(job)
        hydrated["source"] = "company_candidate"
        hydrated["link_status"] = hydrated.get("link_status") or "needs_verification"
        return hydrated

    requested_count = max(1, max_jobs)
    lookup_location = (locations or ["武汉"])[0]
    selected: List[Dict[str, Any]] = []
    seen: set[tuple[str, str]] = set()

    def add_unique(source_jobs: List[Dict[str, Any]]) -> None:
        for job in source_jobs:
            normalized = normalize_candidate_source(job)
            identity = (
                _normalize_search_text(_job_title(normalized)),
                _normalize_search_text(_company_name(normalized)),
            )
            if not all(identity) or identity in seen:
                continue
            if not JobDatabase.has_valid_job_identity(_job_title(normalized), _company_name(normalized)):
                continue
            seen.add(identity)
            selected.append(attach_company_map_url(normalized, lookup_location))
            if len(selected) >= requested_count:
                return

    observed_jobs = search_observed_real_company_jobs(
        keywords=keywords,
        locations=locations,
        max_jobs=max(requested_count * 4, 20),
        candidate_profile=candidate_profile,
        job_type=job_type,
        experience_level=experience_level,
        search_plan=[],
    )
    add_unique(_filter_jobs_for_specific_location(observed_jobs, locations))
    if len(selected) < requested_count:
        # Keep exact district matches first, then fill from the same role/city
        # pool so a requested count of 5 does not collapse to 0 or 1.
        add_unique(observed_jobs)

    if len(selected) < requested_count:
        smart_jobs = build_smart_search_jobs(
            keywords=keywords,
            locations=locations,
            max_jobs=max(requested_count * 3, 15),
            candidate_profile=candidate_profile,
            job_type=job_type,
            experience_level=experience_level,
            search_plan=[],
        )
        add_unique(smart_jobs)

    return _enrich_fast_candidates_with_analysis(
        selected[:requested_count],
        candidate_profile,
    )


# Define API endpoints
@app.get("/")
async def root():
    return {"message": "JobSearch API is running. Access /docs for API documentation."}


@app.get("/health")
async def health_check():
    """Lightweight unauthenticated health check for deployment probes."""
    return {
        "status": "ok",
        "service": "jobsearch-agent",
        "environment": ENVIRONMENT,
    }


@app.get("/imports/boss-collector.js")
async def boss_collector_script():
    """Return a page-side collector that imports visible BOSS job cards."""
    script = r"""
(() => {
  const API_BASE = window.JOB_AGENT_API_BASE || "http://127.0.0.1:8010";
  const text = (root, selector) => root.querySelector(selector)?.textContent?.trim() || "";
  const href = (root, selector) => root.querySelector(selector)?.href || "";
  const abs = (url) => {
    try {
      return url ? new URL(url, location.href).href : "";
    } catch {
      return url || "";
    }
  };
  const uniq = (items) => [...new Set(items.filter(Boolean).map((item) => item.trim()))];
  const readTags = (card) =>
    uniq(Array.from(card.querySelectorAll(".tag-list li, .job-card-footer li, .job-tags span, .info-desc")).map((node) => node.textContent || ""));
  const cards = Array.from(document.querySelectorAll(".job-card-wrapper, .job-list-box li, .job-primary, .job-card-body, [data-jobid]"));
  const jobs = cards
    .map((card) => {
      const name = text(card, ".job-name") || text(card, ".job-title") || text(card, ".job-card-left a") || text(card, "a[href*='job_detail']");
      const company = text(card, ".company-name") || text(card, ".boss-name") || text(card, ".company-text") || text(card, "a[href*='gongsi']");
      const location = text(card, ".job-area") || text(card, ".job-location") || text(card, ".company-location");
      const salary = text(card, ".salary") || text(card, ".red") || text(card, ".job-limit .red");
      const link = abs(href(card, "a[href*='job_detail']"));
      const detail = text(card, ".job-card-footer") || text(card, ".info-desc") || text(card, ".job-limit");
      return {
        name,
        company,
        location,
        salary,
        link,
        detail,
        tags: readTags(card),
        source_page: location.href,
        collected_at: new Date().toISOString()
      };
    })
    .filter((job, index, all) => {
      if (!job.name || !job.company) return false;
      return all.findIndex((item) => item.name === job.name && item.company === job.company) === index;
    });
  if (!jobs.length) {
    alert("没有采集到岗位。请确认当前页面是 BOSS 岗位列表页，并且岗位卡片已经加载出来。");
    return;
  }
  const title = document.querySelector("input[name='query'], .ipt-search")?.value || document.title || "BOSS 页面采集";
  const city = text(document, ".city-label") || text(document, ".job-area") || "";
  fetch(`${API_BASE}/imports/jobs`, {
    method: "POST",
    headers: { "Content-Type": "application/json" },
    body: JSON.stringify({
      source: "boss_page_collector",
      keywords: title,
      locations: city ? [city] : [],
      max_jobs: jobs.length,
      jobs
    })
  })
    .then(async (response) => {
      if (!response.ok) throw new Error(await response.text());
      return response.json();
    })
    .then((data) => alert(`已导入 ${jobs.length} 个真实岗位：${data.search_id}`))
    .catch((error) => {
      console.error(error);
      alert(`导入失败：${error.message || error}`);
    });
})();
""".strip()
    return Response(content=script, media_type="application/javascript; charset=utf-8")


@app.post("/imports/jobs")
async def import_jobs(request: ImportedJobsRequest):
    """Import real job cards collected from an already-open recruiting page."""
    if not request.jobs:
        raise HTTPException(status_code=400, detail="jobs is required")
    try:
        result = save_imported_jobs_as_search(request)
        if str(request.source or "").startswith("boss_monitor"):
            await manager.broadcast(
                {
                    "type": "boss_monitor_update",
                    "data": {
                        "type": "boss_monitor_update",
                        "search_id": result["search_id"],
                        "status": "completed",
                        "job_count": result["job_count"],
                        "updated_at": boss_monitor_state.get("last_run_at"),
                    },
                }
            )
        return result
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/boss/monitor/status")
async def get_boss_monitor_status():
    return boss_monitor_snapshot()


@app.post("/boss/monitor/start", dependencies=[Depends(verify_api_key)])
async def start_boss_monitor(request: BossMonitorRequest):
    """Start or replace the background BOSS monitor."""
    global boss_monitor_task
    config = request.dict()
    config["scrapers"] = ["boss"]

    if boss_monitor_task and not boss_monitor_task.done():
        boss_monitor_state["running"] = False
        boss_monitor_task.cancel()

    boss_monitor_state.update(
        {
            "running": True,
            "keywords": config["keywords"],
            "locations": config["locations"],
            "job_type": config["job_type"],
            "experience_level": config["experience_level"],
            "max_jobs": config["max_jobs"],
            "scrapers": config["scrapers"],
            "interval_seconds": config["interval_seconds"],
            "last_error": None,
        }
    )
    boss_monitor_task = asyncio.create_task(boss_monitor_loop(config))
    await manager.broadcast({"type": "boss_monitor_status", "data": boss_monitor_snapshot()})
    return boss_monitor_snapshot()


@app.post("/boss/monitor/stop", dependencies=[Depends(verify_api_key)])
async def stop_boss_monitor():
    """Stop the background BOSS monitor."""
    global boss_monitor_task
    boss_monitor_state["running"] = False
    if boss_monitor_task and not boss_monitor_task.done():
        boss_monitor_task.cancel()
    await manager.broadcast({"type": "boss_monitor_status", "data": boss_monitor_snapshot()})
    return boss_monitor_snapshot()


@app.post("/search", dependencies=[Depends(verify_api_key)])
async def search_jobs(request: JobSearchRequest, background_tasks: BackgroundTasks):
    """
    Search for jobs based on provided criteria.
    Returns a job ID that can be used to fetch results.
    """
    try:
        # Generate unique ID for this search
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        search_id = f"job_search_{timestamp}"

        if "boss" in [scraper.lower() for scraper in request.scrapers] and not is_boss_automation_enabled():
            results = build_fast_boss_candidates(
                keywords=request.keywords,
                locations=request.locations,
                max_jobs=request.max_jobs,
                candidate_profile=request.candidate_profile,
                job_type=request.job_type,
                experience_level=request.experience_level,
            )
            result_file = os.path.join(output_dir, f"{search_id}.json")
            with open(result_file, "w", encoding="utf-8") as f:
                json.dump(results, f, indent=2, ensure_ascii=False)

            add_search_to_history(
                search_id=search_id,
                keywords=request.keywords,
                locations=request.locations,
                job_type=request.job_type,
                experience_level=request.experience_level,
                scrapers=request.scrapers,
                max_jobs=request.max_jobs,
                status="completed",
            )
            update_search_status(search_id, status="completed", job_count=len(results))
            return {"search_id": search_id, "status": "Job search completed"}

        # Run job search in background when live platform automation is enabled.
        background_tasks.add_task(
            _run_job_search,
            search_id=search_id,
            keywords=request.keywords,
            locations=request.locations,
            job_type=request.job_type,
            experience_level=request.experience_level,
            max_jobs=request.max_jobs,
            scrapers=request.scrapers,
            candidate_profile=request.candidate_profile,
        )

        # Add search to history
        add_search_to_history(
            search_id=search_id,
            keywords=request.keywords,
            locations=request.locations,
            job_type=request.job_type,
            experience_level=request.experience_level,
            scrapers=request.scrapers,
            max_jobs=request.max_jobs,
            status="started",
        )

        return {"search_id": search_id, "status": "Job search started"}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


async def _run_job_search(
    search_id: str,
    keywords: str,
    locations: List[str],
    job_type: str,
    experience_level: str,
    max_jobs: int,
    scrapers: List[str],
    candidate_profile: Optional[str] = None,
):
    """Background task to run job search"""
    try:
        if "boss" in [scraper.lower() for scraper in scrapers]:
            loop = asyncio.get_event_loop()
            results = await loop.run_in_executor(
                None,
                run_boss_deepseek_search,
                keywords,
                locations,
                max_jobs,
                candidate_profile,
                job_type,
                experience_level,
            )
            result_file = os.path.join(output_dir, f"{search_id}.json")
            with open(result_file, "w", encoding="utf-8") as f:
                json.dump(results, f, indent=2, ensure_ascii=False)
            update_search_status(search_id, status="completed", job_count=len(results))
            return

        # Use synchronous version that works reliably on Windows
        # Run in thread pool to avoid blocking the event loop
        loop = asyncio.get_event_loop()

        output_file = await loop.run_in_executor(
            None,
            run_job_search,  # Use sync version
            keywords,
            locations,
            job_type,
            experience_level,
            max_jobs,
            scrapers,
        )

        # Create results file with search_id
        result_file = os.path.join(output_dir, f"{search_id}.json")

        # If output_file is None (database-only mode), create an empty results file
        if output_file is None:
            results = load_search_results_from_database(
                keywords=keywords,
                locations=locations,
                max_jobs=max_jobs,
            )
            with open(result_file, "w", encoding="utf-8") as f:
                json.dump(results, f, indent=2, ensure_ascii=False)
            job_count = len(results)
        else:
            # Copy the output to the result file
            with open(output_file, "r", encoding="utf-8") as src:
                content = src.read()
                with open(result_file, "w", encoding="utf-8") as dst:
                    dst.write(content)

            # Get actual job count from results
            try:
                results = json.loads(content)
                job_count = len(results) if isinstance(results, list) else 0
            except:
                job_count = 0

        # Update search status in history
        update_search_status(search_id, status="completed", job_count=job_count)

    except Exception as e:
        # Log the error
        error_file = os.path.join(output_dir, f"{search_id}_error.txt")
        with open(error_file, "w", encoding="utf-8") as f:
            f.write(str(e))

        # Create empty results file to prevent returning old database jobs
        result_file = os.path.join(output_dir, f"{search_id}.json")
        with open(result_file, "w", encoding="utf-8") as f:
            json.dump([], f)

        # Update search status to error
        update_search_status(search_id, status="error")


# NOTE: Parse endpoint - Will be tested later with agent functionality
# @app.post("/parse")
# async def parse_job_posting(request: JobParseRequest):
#     """Parse job details from text, file content, or URL"""
#     try:
#         # Check if we have required input data
#         if not request.text and not request.file_content and not request.url:
#             raise HTTPException(
#                 status_code=400,
#                 detail="Either text, file_content, or url must be provided",
#             )
#
#         # If URL is provided and extract_webpage is True, set the input type to URL
#         if request.url and request.extract_webpage:
#             # Import web scraping libraries (install if needed)
#             try:
#                 import requests
#                 from bs4 import BeautifulSoup
#             except ImportError:
#                 raise HTTPException(
#                     status_code=500,
#                     detail="Web scraping libraries not installed. Install requests and beautifulsoup4.",
#                 )
#
#             try:
#                 # Fetch the webpage content
#                 headers = {
#                     "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36"
#                 }
#                 response = requests.get(request.url, headers=headers, timeout=10)
#                 response.raise_for_status()
#
#                 # Parse HTML content
#                 soup = BeautifulSoup(response.text, "html.parser")
#
#                 # Extract the main content (this is a simple extraction, can be improved)
#                 job_content = ""
#
#                 # Try to find job description container
#                 job_description = soup.find(
#                     "div", class_=["job-description", "description", "jobDescription"]
#                 )
#                 if job_description:
#                     job_content = job_description.get_text(separator="\n")
#                 else:
#                     # If no specific job container found, extract main content
#                     main_content = (
#                         soup.find("main") or soup.find("article") or soup.find("body")
#                     )
#                     job_content = main_content.get_text(separator="\n")
#
#                 # Add the URL as a source
#                 job_content += f"\n\nSource URL: {request.url}"
#
#                 # Call the job parser agent with the extracted content
#                 parsed_data = call_job_parsr_agent(job_content)
#
#                 # Return parsed data as JSON
#                 return json.loads(parsed_data)
#
#             except requests.RequestException as e:
#                 raise HTTPException(
#                     status_code=500, detail=f"Failed to fetch the webpage: {str(e)}"
#                 )
#         else:
#             # Use text if provided, otherwise use file_content
#             text_to_parse = request.text if request.text else request.file_content
#
#             # Call the job parser agent
#             parsed_data = call_job_parsr_agent(text_to_parse)
#
#             # Return parsed data as JSON
#             return json.loads(parsed_data)
#     except Exception as e:
#         raise HTTPException(status_code=500, detail=str(e))


# NOTE: Process endpoint - Will be tested later with agent functionality
# @app.post("/process")
# async def process_job(request: JobProcessRequest, background_tasks: BackgroundTasks):
#     """
#     Process a job posting to generate CV and/or cover letter
#     Returns a process_id that can be used to fetch results
#     """
#     try:
#         # Generate unique ID for this process
#         timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
#         process_id = f"job_process_{timestamp}"
#
#         # Run process in background
#         background_tasks.add_task(
#             _run_job_process,
#             process_id=process_id,
#             job_posting=request.job_posting,
#             generate_cv=request.generate_cv,
#             generate_cover_letter=request.generate_cover_letter,
#         )
#
#         return {"process_id": process_id, "status": "Job processing started"}
#     except Exception as e:
#         raise HTTPException(status_code=500, detail=str(e))


# NOTE: Background job processing function - Will be tested later with agent functionality
# async def _run_job_process(
#     process_id: str,
#     job_posting: Dict[str, Any],
#     generate_cv: bool,
#     generate_cover_letter: bool,
# ):
#     """Background task to process job posting"""
#     try:
#         # Convert job posting to json string
#         job_details_str = json.dumps(job_posting)
#
#         # Initialize document database
#         from src.utils.document_database import DocumentStorage
#
#         results = {
#             "status": "completed",
#             "documents": {},
#             "cv_content": None,
#             "cover_letter_content": None,
#         }
#
#         # Create folder only for metadata (no more document files)
#         company = job_posting.get("company_name", "Unknown")
#         job_title = job_posting.get("job_title", "Job")
#         company_slug = slugify(company)
#         job_title_slug = slugify(job_title)
#         folder_name = os.path.join(output_dir, slugify(f"{company}_{job_title}"))
#         os.makedirs(folder_name, exist_ok=True)
#
#         # Save job metadata only
#         metadata_path = os.path.join(
#             folder_name, f"{company_slug}_{job_title_slug}_metadata.json"
#         )
#         with open(metadata_path, "w", encoding="utf-8") as f:
#             json.dump(job_posting, f, indent=2)
#
#         # Generate CV if requested
#         if generate_cv:
#             try:
#                 cv_text, state_json, cv_path = call_cv_agent(job_details_str)
#
#                 # Store CV in database
#                 cv_id = DocumentStorage.store_cv(
#                     content=cv_text,
#                     job_posting=job_posting,
#                     process_id=process_id,
#                     state_json=state_json,
#                     template_used="cv_template.txt",  # Could be made configurable
#                 )
#
#                 results["documents"]["cv_id"] = cv_id
#                 results["cv_content"] = cv_text
#
#                 print(f"✅ CV generated and stored in database (ID: {cv_id})")
#
#             except Exception as e:
#                 print(f"❌ Error generating CV: {e}")
#                 results["cv_error"] = str(e)
#
#         # Generate cover letter if requested
#         if generate_cover_letter:
#             try:
#                 cover_letter_text, cl_state_json, cl_path = call_cover_letter_agent(
#                     job_details_str
#                 )
#
#                 # Store cover letter in database
#                 cl_id = DocumentStorage.store_cover_letter(
#                     content=cover_letter_text,
#                     job_posting=job_posting,
#                     process_id=process_id,
#                     state_json=cl_state_json,
#                     template_used="cover_letter_template.txt",  # Could be made configurable
#                 )
#
#                 results["documents"]["cover_letter_id"] = cl_id
#                 results["cover_letter_content"] = cover_letter_text
#
#                 print(f"✅ Cover letter generated and stored in database (ID: {cl_id})")
#
#             except Exception as e:
#                 print(f"❌ Error generating cover letter: {e}")
#                 results["cover_letter_error"] = str(e)
#
#         # Write results file
#         with open(
#             os.path.join(output_dir, f"{process_id}.json"), "w", encoding="utf-8"
#         ) as f:
#             json.dump(results, f, indent=2)
#
#     except Exception as e:
#         # Log the error
#         print(f"❌ Error in job processing: {e}")
#         with open(
#             os.path.join(output_dir, f"{process_id}_error.txt"), "w", encoding="utf-8"
#         ) as f:
#             f.write(str(e))


# NOTE: Get process results endpoint - Will be tested later with agent functionality
# @app.get("/process/{process_id}")
# async def get_process_results(process_id: str):
#     """Get the results of a job processing task by ID"""
#     result_file = os.path.join(output_dir, f"{process_id}.json")
#     error_file = os.path.join(output_dir, f"{process_id}_error.txt")
#
#     if os.path.exists(result_file):
#         with open(result_file, "r", encoding="utf-8") as f:
#             return json.load(f)
#     elif os.path.exists(error_file):
#         with open(error_file, "r", encoding="utf-8") as f:
#             error = f.read()
#         raise HTTPException(status_code=500, detail=error)
#     else:
#         return {"status": "in_progress", "message": "Job processing is still running"}


# WebSocket endpoint for real-time interaction with agents
@app.websocket("/ws")
async def websocket_endpoint(websocket: WebSocket):
    """WebSocket endpoint for real-time communication"""
    await manager.connect(websocket)
    try:
        while True:  # Receive message from client
            data = await websocket.receive_text()
            message = json.loads(data)

            action = message.get("action")

            if action == "search":
                await handle_ws_search(websocket, message)
            else:
                await manager.send_error(websocket, f"Unknown action: {action}")

    except Exception as e:
        print(f"WebSocket error: {e}")
    finally:
        manager.disconnect(websocket)


async def handle_ws_search(websocket: WebSocket, message_data: dict):
    """Handle job search request via WebSocket with real-time progress"""
    try:
        data = message_data.get("data", {})

        # Extract search parameters
        keywords = data.get("keywords", "")
        locations = data.get("locations", ["Remote"])
        job_type = data.get("job_type", "full-time")
        experience_level = data.get("experience_level", "mid-level")
        max_jobs = data.get("max_jobs", 3)
        scrapers = data.get("scrapers", ["linkedin"])
        candidate_profile = data.get("candidate_profile")

        if not keywords:
            await manager.send_error(websocket, "Keywords are required")
            return

        # Generate unique ID for this search
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        search_id = f"job_search_{timestamp}"

        await manager.send_progress(websocket, f"Starting job search: {keywords}")

        # Check for similar searches first
        await manager.send_progress(
            websocket, "Checking for similar recent searches..."
        )
        similar_searches = find_similar_searches(
            keywords=keywords,
            locations=locations,
            job_type=job_type,
            experience_level=experience_level,
            scrapers=scrapers,
        )

        # Check if we have recent exact matches
        now = datetime.now()
        recent_exact = []
        for search in similar_searches["exact"]:
            search_time = datetime.fromisoformat(search["timestamp"])
            if (now - search_time).total_seconds() < 86400:  # 24 hours
                result_file = os.path.join(output_dir, f"{search['search_id']}.json")
                if os.path.exists(result_file):
                    recent_exact.append(search)

        if recent_exact:
            await manager.send_progress(
                websocket, "Found recent similar search results"
            )
            # Send the similar searches for user decision
            await websocket.send_json(
                {
                    "type": "similar_found",
                    "data": {
                        "exact_matches": recent_exact,
                        "similar_searches": similar_searches["similar"][:3],
                        "search_id": search_id,
                    },
                }
            )
            return

        # No recent matches, proceed with new search
        await manager.send_progress(
            websocket, "No recent similar searches found, starting new search..."
        )

        # Add search to history
        add_search_to_history(
            search_id=search_id,
            keywords=keywords,
            locations=locations,
            job_type=job_type,
            experience_level=experience_level,
            scrapers=scrapers,
            max_jobs=max_jobs,
            status="started",
        )

        # Run job search with progress updates
        await _run_job_search_with_websocket(
            websocket,
            search_id=search_id,
            keywords=keywords,
            locations=locations,
            job_type=job_type,
            experience_level=experience_level,
            max_jobs=max_jobs,
            scrapers=scrapers,
            candidate_profile=candidate_profile,
        )

    except Exception as e:
        await manager.send_error(websocket, str(e))


async def _run_job_search_with_websocket(
    websocket: WebSocket,
    search_id: str,
    keywords: str,
    locations: List[str],
    job_type: str,
    experience_level: str,
    max_jobs: int,
    scrapers: List[str],
    candidate_profile: Optional[str] = None,
):
    """Background task to run job search with WebSocket progress updates"""
    try:
        await manager.send_progress(
            websocket, f"Initializing scrapers: {', '.join(scrapers)}"
        )

        if "boss" in [scraper.lower() for scraper in scrapers]:
            await manager.send_progress(websocket, "正在搜索 BOSS 直聘并调用 DeepSeek 分析...")
            loop = asyncio.get_event_loop()
            results = await loop.run_in_executor(
                None,
                run_boss_deepseek_search,
                keywords,
                locations,
                max_jobs,
                candidate_profile,
            )
            result_file = os.path.join(output_dir, f"{search_id}.json")
            with open(result_file, "w", encoding="utf-8") as f:
                json.dump(results, f, indent=2, ensure_ascii=False)
            update_search_status(search_id, status="completed", job_count=len(results))
            await manager.send_result(
                websocket,
                {
                    "search_id": search_id,
                    "jobs": results,
                    "job_count": len(results),
                    "status": "completed",
                },
            )
            return

        # Initialize progress tracking
        progress_steps = [
            "Setting up search parameters",
            "Starting web scrapers",
            "Scraping job listings",
            "Processing job data",
            "Filtering and cleaning results",
            "Finalizing search results",
        ]

        for i, step in enumerate(progress_steps):
            await manager.send_progress(
                websocket, f"Step {i + 1}/{len(progress_steps)}: {step}"
            )
            await asyncio.sleep(0.5)  # Small delay for realistic progress

        # Run job search pipeline using sync version in thread pool
        await manager.send_progress(websocket, "Running job search pipeline...")

        loop = asyncio.get_event_loop()
        output_file = await loop.run_in_executor(
            None,
            run_job_search,  # Use sync version
            keywords,
            locations,
            job_type,
            experience_level,
            max_jobs,
            scrapers,
        )

        # Create results file with search_id
        result_file = os.path.join(output_dir, f"{search_id}.json")

        # Handle both database-only and file output modes
        if output_file is None:
            results = load_search_results_from_database(
                keywords=keywords,
                locations=locations,
                max_jobs=max_jobs,
            )
            with open(result_file, "w", encoding="utf-8") as f:
                json.dump(results, f, indent=2, ensure_ascii=False)
            job_count = len(results)
        else:
            # Copy the output to the result file
            with open(output_file, "r", encoding="utf-8") as src:
                content = src.read()
                with open(result_file, "w", encoding="utf-8") as dst:
                    dst.write(content)

            # Get actual job count from results
            try:
                results = json.loads(content)
                job_count = len(results) if isinstance(results, list) else 0
            except:
                job_count = 0
                results = []

        await manager.send_progress(
            websocket, f"Search completed! Found {job_count} jobs"
        )

        # Update search status in history
        update_search_status(search_id, status="completed", job_count=job_count)

        # Send final results
        await manager.send_result(
            websocket,
            {
                "search_id": search_id,
                "jobs": results,
                "job_count": job_count,
                "status": "completed",
            },
        )

    except Exception as e:
        # Log the error
        error_file = os.path.join(output_dir, f"{search_id}_error.txt")
        with open(error_file, "w", encoding="utf-8") as f:
            f.write(str(e))

        # Create empty results file to prevent returning old database jobs
        result_file = os.path.join(output_dir, f"{search_id}.json")
        with open(result_file, "w", encoding="utf-8") as f:
            json.dump([], f)

        # Update search status to error
        update_search_status(search_id, status="error")
        await manager.send_error(websocket, f"Search failed: {str(e)}")


@app.post("/search/check")
async def check_similar_searches(request: JobSearchRequest):
    """Check for existing/similar searches before starting a new one"""
    try:
        similar_searches = find_similar_searches(
            keywords=request.keywords,
            locations=request.locations,
            job_type=request.job_type,
            experience_level=request.experience_level,
            scrapers=request.scrapers,
        )

        # Filter out old searches (older than 24 hours) for exact matches
        now = datetime.now()
        recent_exact = []
        for search in similar_searches["exact"]:
            search_time = datetime.fromisoformat(search["timestamp"])
            if (now - search_time).total_seconds() < 86400:  # 24 hours
                # Check if results file exists
                result_file = os.path.join(output_dir, f"{search['search_id']}.json")
                if os.path.exists(result_file):
                    # Add job count if available
                    try:
                        with open(result_file, "r", encoding="utf-8") as f:
                            results = json.load(f)
                            search["job_count"] = (
                                len(results) if isinstance(results, list) else 0
                            )
                    except:
                        search["job_count"] = 0
                    recent_exact.append(search)

        # Get recent similar searches (last 7 days)
        recent_similar = []
        for search in similar_searches["similar"][:5]:  # Limit to 5 most recent
            search_time = datetime.fromisoformat(search["timestamp"])
            if (now - search_time).total_seconds() < 604800:  # 7 days
                result_file = os.path.join(output_dir, f"{search['search_id']}.json")
                if os.path.exists(result_file):
                    try:
                        with open(result_file, "r", encoding="utf-8") as f:
                            results = json.load(f)
                            search["job_count"] = (
                                len(results) if isinstance(results, list) else 0
                            )
                    except:
                        search["job_count"] = 0
                    recent_similar.append(search)

        return {
            "exact_matches": recent_exact,
            "similar_searches": recent_similar,
            "has_recent_exact": len(recent_exact) > 0,
        }

    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/boss/analyze-text")
async def analyze_boss_text(request: BossTextAnalyzeRequest):
    """Analyze pasted BOSS/JD text with the same DeepSeek Agent pipeline."""
    if not request.job_text.strip():
        raise HTTPException(status_code=400, detail="job_text is required")

    try:
        job = parse_boss_job_text(
            job_text=request.job_text,
            source_url=request.source_url,
        )
        job = attach_company_map_url(job)
        job["ai_analysis"] = DeepSeekClient().analyze_job(
            job,
            candidate_profile=request.candidate_profile,
        )
        return {"success": True, "data": job}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/career/analyze")
async def analyze_career_fit(request: CareerAnalyzeRequest):
    """Recommend suitable jobs and learning gaps from resume-first career criteria."""
    if not request.candidate_profile.strip():
        raise HTTPException(status_code=400, detail="candidate_profile is required")
    if not request.target_role.strip():
        raise HTTPException(status_code=400, detail="target_role is required")

    try:
        recommended_jobs = build_fast_boss_candidates(
            keywords=request.target_role,
            locations=[request.location],
            max_jobs=request.max_recommendations,
            candidate_profile=request.candidate_profile,
            job_type=request.job_type,
            experience_level=request.experience_level,
            strict_location=True,
        )
        market_trends = build_job_trends(
            recommended_jobs,
            target_role=request.target_role,
            location=request.location,
        )
        career_analysis = DeepSeekClient().analyze_career_fit(
            candidate_profile=request.candidate_profile,
            target_role=request.target_role,
            location=request.location,
            job_type=request.job_type,
            experience_level=request.experience_level,
            requirement_text=request.requirement_text,
            recommended_jobs=recommended_jobs,
        )
        career_analysis["market_trends"] = market_trends
        return {
            "success": True,
            "data": {
                "target_role": request.target_role,
                "location": request.location,
                "job_type": request.job_type,
                "experience_level": request.experience_level,
                "recommended_jobs": recommended_jobs,
                "market_trends": market_trends,
                "career_analysis": career_analysis,
            },
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/resume/extract")
async def extract_resume(request: ResumeExtractRequest):
    """Extract plain text from a resume file so the Agent can use a real profile."""
    try:
        content = base64.b64decode(request.content_base64)
        text = extract_resume_text(request.file_name, content)
        if not text:
            raise HTTPException(status_code=400, detail="简历里没有提取到文本内容。")
        return {"success": True, "data": {"text": text[:6000]}}
    except HTTPException:
        raise
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/search/history")
async def get_search_history(limit: int = 20):
    """Get recent search history"""
    try:
        history = load_search_history()

        # Add job counts and filter valid searches
        valid_history = []
        for search in history[:limit]:
            result_file = os.path.join(output_dir, f"{search['search_id']}.json")
            search["job_count"] = search.get("job_count", 0) or 0
            search["has_results"] = False

            if os.path.exists(result_file):
                try:
                    with open(result_file, "r", encoding="utf-8") as f:
                        results = json.load(f)
                    if isinstance(results, list):
                        filtered_results = _filter_jobs_for_search(
                            jobs=results,
                            keywords=search.get("keywords", ""),
                            locations=search.get("locations", []),
                            max_jobs=search.get("max_jobs"),
                        )
                        if not filtered_results and search.get("status") == "completed":
                            filtered_results = regenerate_visible_search_results(search)
                        search["job_count"] = len(filtered_results)
                        search["has_results"] = len(filtered_results) > 0
                    else:
                        search["job_count"] = 0
                        search["has_results"] = False
                except:
                    search["job_count"] = 0
                    search["has_results"] = False

            valid_history.append(search)

        return {"searches": valid_history}

    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.delete("/search/history")
async def clear_search_history():
    """Delete all recent search cards and their result files."""
    try:
        history = load_search_history()
        removed_files = 0
        for search in history:
            removed_files += delete_search_artifacts(str(search.get("search_id", "")))
        save_search_history([])
        return {
            "success": True,
            "deleted": len(history),
            "removed_files": removed_files,
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.delete("/search/history/{search_id}")
async def delete_search_history_item(search_id: str):
    """Delete one recent search card and its result files."""
    try:
        history = load_search_history()
        kept = [search for search in history if search.get("search_id") != search_id]
        if len(kept) == len(history):
            raise HTTPException(status_code=404, detail="Search history item not found")
        removed_files = delete_search_artifacts(search_id)
        save_search_history(kept)
        return {
            "success": True,
            "deleted": 1,
            "removed_files": removed_files,
        }
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# IMPORTANT: Specific routes must come before parametrized routes in FastAPI
# Otherwise /search/history would match /search/{search_id} with search_id="history"
@app.get("/search/{search_id}")
async def get_search_results(search_id: str):
    """Get the results of a job search by ID"""
    result_file = os.path.join(output_dir, f"{search_id}.json")
    error_file = os.path.join(output_dir, f"{search_id}_error.txt")

    if os.path.exists(result_file):
        with open(result_file, "r", encoding="utf-8") as f:
            results = json.load(f)
        search = _find_search_history_entry(search_id)
        if isinstance(results, list) and search:
            filtered_results = _filter_jobs_for_search(
                jobs=results,
                keywords=search.get("keywords", ""),
                locations=search.get("locations", []),
                max_jobs=search.get("max_jobs"),
            )
            if not filtered_results and search.get("status") == "completed":
                return regenerate_visible_search_results(search)
            return filtered_results
        return results
    elif os.path.exists(error_file):
        with open(error_file, "r", encoding="utf-8") as f:
            error = f.read()
        raise HTTPException(status_code=500, detail=error)
    else:
        return {"status": "in_progress", "message": "Job search is still running"}


# ============================================================================
# DATABASE ENDPOINTS
# ============================================================================


@app.get("/jobs/stats")
async def get_job_stats():
    """Get database statistics"""
    try:
        db = JobDatabase()
        stats = db.get_stats()
        db.close()
        return {"success": True, "data": stats}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/jobs")
async def get_jobs(limit: int = 100, offset: int = 0):
    """Get jobs from database with pagination"""
    try:
        db = JobDatabase()
        jobs = db.get_jobs(limit=limit, offset=offset)
        db.close()
        return {"success": True, "data": jobs, "count": len(jobs)}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/jobs/search")
async def search_jobs_db(
    keyword: str = None, company: str = None, location: str = None
):
    """Search jobs in database"""
    try:
        db = JobDatabase()
        jobs = db.search_jobs(keyword=keyword, company=company, location=location)
        db.close()
        return {"success": True, "data": jobs, "count": len(jobs)}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/jobs/{job_id}")
async def get_job(job_id: int):
    """Get a specific job by ID"""
    try:
        db = JobDatabase()
        job = db.get_job(job_id)
        db.close()
        if not job:
            raise HTTPException(status_code=404, detail="Job not found")

        return {"success": True, "data": job}
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/jobs/migrate")
async def migrate_jobs():
    """Migrate existing JSON files to database"""
    try:
        import glob

        # Find all JSON files in jobs directory
        json_files = glob.glob("jobs/*.json")
        if not json_files:
            return {
                "success": True,
                "message": "No JSON files found to migrate",
                "migrated": 0,
            }

        db = JobDatabase()
        migrated_count = db.migrate_from_json(json_files)
        stats = db.get_stats()
        db.close()

        return {
            "success": True,
            "message": f"Migration completed. {migrated_count} jobs migrated.",
            "migrated": migrated_count,
            "total_in_db": stats["total_jobs"],
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/interview-logs")
async def create_interview_log(request: InterviewLogRequest):
    """Store a new interview log."""
    try:
        db = get_job_database()
        try:
            log_id = db.add_interview_log(
                job_id=request.job_id,
                job_title=request.job_title,
                company_name=request.company_name,
                interview_date=request.interview_date,
                outcome=request.outcome,
                failure_reason=request.failure_reason,
                notes=request.notes,
                next_action=request.next_action,
            )
            log = db.get_interview_logs(limit=1, offset=0, company_name=request.company_name)
            created_log = next((item for item in log if item["id"] == log_id), None)
            return {"success": True, "data": created_log or {"id": log_id}}
        finally:
            db.close()
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/interview-logs")
async def get_interview_logs(
    limit: int = 100,
    offset: int = 0,
    company_name: str = None,
    outcome: str = None,
    job_id: int = None,
):
    """List interview logs with optional filters."""
    try:
        db = get_job_database()
        try:
            logs = db.get_interview_logs(
                limit=limit,
                offset=offset,
                company_name=company_name,
                outcome=outcome,
                job_id=job_id,
            )
            return {"success": True, "data": logs, "count": len(logs)}
        finally:
            db.close()
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/interview-logs/stats")
async def get_interview_log_stats():
    """Return interview log summary statistics."""
    try:
        db = get_job_database()
        try:
            stats = db.get_interview_log_stats()
            return {"success": True, "data": stats}
        finally:
            db.close()
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# ============================================================================
# DOCUMENT DATABASE ENDPOINTS
# ============================================================================
# NOTE: Document database endpoints - Will be tested later with agent functionality

# @app.get("/documents/stats")
# async def get_document_stats():
#     """Get document database statistics"""
#     try:
#         from src.utils.document_database import DocumentDatabase
#
#         db = DocumentDatabase()
#         stats = db.get_document_stats()
#         db.close()
#         return {"success": True, "data": stats}
#     except Exception as e:
#         raise HTTPException(status_code=500, detail=str(e))
#
#
# @app.get("/documents")
# async def get_documents(document_type: str = None, limit: int = 50):
#     """Get recent documents with optional filtering by type"""
#     try:
#         from src.utils.document_database import DocumentDatabase
#
#         db = DocumentDatabase()
#         documents = db.get_recent_documents(document_type=document_type, limit=limit)
#         db.close()
#         return {"success": True, "data": documents, "count": len(documents)}
#     except Exception as e:
#         raise HTTPException(status_code=500, detail=str(e))
#
#
# @app.get("/documents/search")
# async def search_documents(
#     keyword: str = None, company: str = None, document_type: str = None
# ):
#     """Search documents by keyword, company, or type"""
#     try:
#         from src.utils.document_database import DocumentDatabase
#
#         db = DocumentDatabase()
#         documents = db.search_documents(
#             keyword=keyword, company=company, document_type=document_type
#         )
#         db.close()
#         return {"success": True, "data": documents, "count": len(documents)}
#     except Exception as e:
#         raise HTTPException(status_code=500, detail=str(e))
#
#
# @app.get("/documents/{document_id}")
# async def get_document(document_id: int):
#     """Get a specific document by ID"""
#     try:
#         from src.utils.document_database import DocumentDatabase
#
#         db = DocumentDatabase()
#         document = db.get_document(document_id)
#         db.close()
#         if not document:
#             raise HTTPException(status_code=404, detail="Document not found")
#         return {"success": True, "data": document}
#     except HTTPException:
#         raise
#     except Exception as e:
#         raise HTTPException(status_code=500, detail=str(e))
#
#
# @app.get("/documents/{document_id}/versions")
# async def get_document_versions(document_id: int):
#     """Get all versions of a document"""
#     try:
#         from src.utils.document_database import DocumentDatabase
#
#         db = DocumentDatabase()
#         versions = db.get_document_versions(document_id)
#         db.close()
#         return {"success": True, "data": versions, "count": len(versions)}
#     except Exception as e:
#         raise HTTPException(status_code=500, detail=str(e))
#
#
# @app.get("/process/{process_id}/documents")
# async def get_process_documents(process_id: str):
#     """Get all documents for a specific process"""
#     try:
#         from src.utils.document_database import DocumentStorage
#
#         documents = DocumentStorage.get_documents_for_process(process_id)
#         return {"success": True, "data": documents}
#     except Exception as e:
#         raise HTTPException(status_code=500, detail=str(e))
#
#
# @app.post("/documents/{document_id}/export")
# async def export_document_to_file(document_id: int, output_dir: str = "output"):
#     """Export a document to a text file"""
#     try:
#         from src.utils.document_database import DocumentDatabase
#
#         db = DocumentDatabase()
#         filepath = db.export_document_to_file(document_id, output_dir)
#         db.close()
#
#         # Return relative path for frontend access
#         relative_path = filepath.replace(os.path.normpath(output_dir), "output")
#         return {
#             "success": True,
#             "filepath": relative_path,
#             "message": f"Document exported to {relative_path}",
#         }
#     except Exception as e:
#         raise HTTPException(status_code=500, detail=str(e))
#
#
# @app.delete("/documents/{document_id}")
# async def delete_document(document_id: int):
#     """Delete a document and all its versions"""
#     try:
#         from src.utils.document_database import DocumentDatabase
#
#         db = DocumentDatabase()
#         success = db.delete_document(document_id)
#         db.close()
#
#         if success:
#             return {
#                 "success": True,
#                 "message": f"Document {document_id} deleted successfully",
#             }
#         else:
#             raise HTTPException(status_code=500, detail="Failed to delete document")
#     except HTTPException:
#         raise
#     except Exception as e:
#         raise HTTPException(status_code=500, detail=str(e))


if __name__ == "__main__":
    # Don't force Windows compatibility here - let the browser manager handle it
    # Run the FastAPI app with uvicorn
    uvicorn.run("main_api:app", host="0.0.0.0", port=8000, reload=True)
