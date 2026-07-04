import os
import base64
import json
from io import BytesIO
import tempfile
import unittest
from pathlib import Path
from unittest.mock import Mock, patch

from fastapi.testclient import TestClient

import main_api
from src.scraper.search.boss_scraper import (
    build_smart_search_jobs,
    build_boss_search_url,
    is_boss_security_verification_url,
    is_boss_automation_enabled,
    get_boss_browser_settings,
    parse_boss_job_cards,
    run_boss_deepseek_search,
    search_curated_agent_jobs,
    search_web_real_company_jobs,
)
from src.utils.deepseek_client import DeepSeekClient


class BossDeepSeekFlowTests(unittest.TestCase):
    def setUp(self):
        self._boss_env_patch = patch.dict(os.environ, {"BOSS_AUTOMATION_ENABLED": "false"}, clear=False)
        self._boss_env_patch.start()
        self.addCleanup(self._boss_env_patch.stop)

    def test_builds_wuhan_boss_search_url(self):
        url = build_boss_search_url("python实习", "武汉")

        self.assertIn("zhipin.com", url)
        self.assertIn("query=python", url)
        self.assertIn("city=101200100", url)

    def test_boss_automation_is_opt_in_to_avoid_browser_jump(self):
        with patch.dict(os.environ, {}, clear=True):
            self.assertFalse(is_boss_automation_enabled())

    def test_parses_boss_job_cards_from_html(self):
        html = """
        <div class="job-card-wrapper">
          <a href="/job_detail/abc123.html">
            <span class="job-name">Python Agent 实习生</span>
          </a>
          <span class="salary">150-200元/天</span>
          <span class="job-area">武汉 洪山区</span>
          <h3 class="company-name">武汉智能科技有限公司</h3>
          <ul class="tag-list"><li>Python</li><li>大模型</li></ul>
        </div>
        """

        jobs = parse_boss_job_cards(html, max_jobs=5)

        self.assertEqual(len(jobs), 1)
        self.assertEqual(jobs[0]["job_title"], "Python Agent 实习生")
        self.assertEqual(jobs[0]["company_name"], "武汉智能科技有限公司")
        self.assertEqual(jobs[0]["job_location"], "武汉 洪山区")
        self.assertEqual(jobs[0]["salary"], "150-200元/天")
        self.assertEqual(jobs[0]["source"], "boss")
        self.assertTrue(jobs[0]["source_url"].startswith("https://www.zhipin.com"))

    def test_deepseek_client_returns_structured_analysis(self):
        session = Mock()
        response = Mock()
        response.raise_for_status.return_value = None
        response.json.return_value = {
            "choices": [
                {
                    "message": {
                        "content": '{"match_score": 82, "summary": "适合投递", "skill_gaps": ["FastAPI"], "resume_tips": ["突出 Python 项目"], "interview_questions": ["讲一下异步"], "recommendation": "建议投递"}'
                    }
                }
            ]
        }
        session.post.return_value = response

        client = DeepSeekClient(api_key="test-key", session=session)
        analysis = client.analyze_job(
            {
                "job_title": "Python 实习生",
                "company_name": "武汉智能科技有限公司",
                "job_description": "负责 Python Agent 开发",
            },
            candidate_profile="大三计科，做过 FastAPI 求职助手项目，Python 基础较弱。",
        )

        self.assertEqual(analysis["match_score"], 82)
        self.assertEqual(analysis["recommendation"], "建议投递")
        self.assertEqual(analysis["provider"], "deepseek")
        session.post.assert_called_once()
        sent_prompt = session.post.call_args.kwargs["json"]["messages"][1]["content"]
        self.assertIn("做过 FastAPI 求职助手项目", sent_prompt)
        self.assertIn("Python 基础较弱", sent_prompt)

    def test_deepseek_prompt_uses_default_profile_when_resume_is_missing(self):
        client = DeepSeekClient(api_key="test-key", session=Mock())

        prompt = client._build_prompt({"job_title": "Python 实习生"}, candidate_profile=None)

        self.assertIn("学生画像", prompt)
        self.assertIn("未上传简历", prompt)

    def test_deepseek_client_can_disable_environment_proxy(self):
        with patch.dict(os.environ, {"DEEPSEEK_DISABLE_PROXY": "true"}, clear=False):
            client = DeepSeekClient(api_key="test-key")

        self.assertFalse(client.session.trust_env)

    def test_api_loads_project_env_path_explicitly(self):
        self.assertTrue(main_api.DOTENV_PATH.endswith(".env"))
        self.assertEqual(Path(main_api.DOTENV_PATH).parent, Path(main_api.__file__).parent)

    def test_boss_search_endpoint_writes_fast_candidate_results(self):
        temp_dir = tempfile.TemporaryDirectory()
        self.addCleanup(temp_dir.cleanup)
        client = TestClient(main_api.app)

        fake_jobs = [
            {
                "job_title": "Python Agent 实习生",
                "company_name": "武汉智能科技有限公司",
                "job_location": "武汉 洪山区",
                "salary": "150-200元/天",
                "source": "boss",
                "source_url": "https://www.zhipin.com/job_detail/abc123.html",
            }
        ]

        with patch("main_api.output_dir", temp_dir.name), patch(
            "main_api.search_history_file",
            os.path.join(temp_dir.name, "search_history.json"),
        ), patch("main_api.is_boss_automation_enabled", return_value=False), patch(
            "main_api.build_fast_boss_candidates", return_value=fake_jobs
        ):
            response = client.post(
                "/search",
                json={
                    "keywords": "python实习",
                    "locations": ["武汉"],
                    "job_type": "internship",
                    "experience_level": "entry-level",
                    "max_jobs": 5,
                    "scrapers": ["boss"],
                },
                headers={"X-API-Key": main_api.API_KEY},
            )

            self.assertEqual(response.status_code, 200)
            search_id = response.json()["search_id"]
            results_response = client.get(f"/search/{search_id}")

        self.assertEqual(results_response.status_code, 200)
        results = results_response.json()
        self.assertEqual(results[0]["company_name"], "武汉智能科技有限公司")
        self.assertEqual(results[0]["source"], "boss")

    def test_search_api_generates_smart_candidates_for_csharp_when_no_observed_match(self):
        temp_dir = tempfile.TemporaryDirectory()
        self.addCleanup(temp_dir.cleanup)
        client = TestClient(main_api.app)
        fake_live_jobs = [
            {
                "source": "web_search",
                "source_url": "https://www.zhipin.com/web/geek/job?query=csharp",
                "link_status": "needs_verification",
                "job_title": "C# .NET开发实习生",
                "company_name": "武汉真实软件有限公司",
                "job_location": "武汉",
                "salary": "150-200元/天",
                "job_description": "C# .NET ASP.NET 实习",
                "tags": ["C#", ".NET", "实习"],
            }
        ]

        with patch("main_api.output_dir", temp_dir.name), patch(
            "main_api.search_history_file",
            os.path.join(temp_dir.name, "search_history.json"),
        ), patch("main_api.is_boss_automation_enabled", return_value=False), patch(
            "main_api.search_web_real_company_jobs", return_value=fake_live_jobs
        ):
            response = client.post(
                "/search",
                json={
                    "keywords": "C# .NET ASP.NET 实习 初级",
                    "locations": ["武汉"],
                    "job_type": "internship",
                    "experience_level": "entry-level",
                    "max_jobs": 3,
                    "scrapers": ["boss"],
                    "candidate_profile": "大三计科，学过 C 语言、Java、MySQL、Vue。",
                },
                headers={"X-API-Key": main_api.API_KEY},
            )

            self.assertEqual(response.status_code, 200)
            search_id = response.json()["search_id"]
            results_response = client.get(f"/search/{search_id}")

        self.assertEqual(results_response.status_code, 200)
        results = results_response.json()
        self.assertEqual(len(results), 1)
        self.assertTrue(any(job["source"] == "web_search" for job in results))
        self.assertTrue(any("C#" in job["job_title"] for job in results))

    def test_blocked_boss_result_is_visible_after_filtering(self):
        blocked = {
            "source": "boss",
            "job_title": "BOSS 自动采集被拦截",
            "company_name": "需要人工验证",
            "job_location": "武汉",
            "ai_analysis": {
                "provider": "deepseek",
                "status": "blocked",
                "summary": "BOSS 触发验证码",
            },
        }

        results = main_api._filter_jobs_for_search(
            jobs=[blocked],
            keywords="python实习",
            locations=["武汉"],
            max_jobs=5,
        )

        self.assertEqual(len(results), 1)
        self.assertEqual(results[0]["ai_analysis"]["status"], "blocked")

    def test_reads_boss_browser_settings_for_visible_persistent_session(self):
        with patch.dict(
            os.environ,
            {"BOSS_HEADLESS": "false", "BOSS_USER_DATA_DIR": "data/boss-browser"},
            clear=False,
        ):
            settings = get_boss_browser_settings()

        self.assertFalse(settings["headless"])
        self.assertTrue(settings["user_data_dir"].replace("\\", "/").endswith("data/boss-browser"))
        self.assertGreaterEqual(settings["manual_check_seconds"], 1)

    def test_detects_boss_security_verification_redirect(self):
        url = (
            "https://www.zhipin.com/web/passport/zp/security.html"
            "?callbackUrl=https%3A%2F%2Fwww.zhipin.com%2Fweb%2Fgeek%2Fjob"
        )

        self.assertTrue(is_boss_security_verification_url(url))

    def test_curated_search_returns_job_names_and_links_without_boss_browser(self):
        jobs = search_curated_agent_jobs("python agent", ["武汉"], 3)

        self.assertGreaterEqual(len(jobs), 1)
        self.assertIn("Python", jobs[0]["job_description"])
        self.assertTrue(jobs[0]["source_url"].startswith("http"))
        self.assertNotIn("/job_detail/", jobs[0]["source_url"])
        self.assertEqual(jobs[0]["link_status"], "needs_verification")
        self.assertNotEqual(jobs[0]["job_title"], "BOSS 自动采集被拦截")

    def test_smart_search_generates_java_jobs_for_java_keyword(self):
        jobs = build_smart_search_jobs(
            "java",
            ["武汉"],
            3,
            candidate_profile="大三计科，学过 Java、MySQL、Vue 和接口自动化。",
        )

        self.assertEqual(len(jobs), 3)
        self.assertTrue(all("Java" in job["job_title"] or "后端" in job["job_title"] for job in jobs))
        self.assertFalse(any("Python" in job["job_title"] or "Agent" in job["job_title"] for job in jobs))
        self.assertTrue(all(job["source"] == "company_candidate" for job in jobs))
        self.assertTrue(all(job["company_name"] != "平台搜索待核验" for job in jobs))
        self.assertTrue(all(job["amap_company_url"].startswith("https://www.amap.com/search") for job in jobs))

    def test_smart_search_generates_dotnet_jobs_for_csharp_keyword(self):
        jobs = build_smart_search_jobs(
            "C# .NET ASP.NET 实习 初级",
            ["武汉"],
            3,
            candidate_profile="大三计科，学过 C 语言、Java、MySQL、Vue。",
        )

        self.assertEqual(len(jobs), 3)
        self.assertTrue(any("C#" in job["job_title"] or ".NET" in job["job_title"] for job in jobs))
        self.assertTrue(all(job["source"] == "company_candidate" for job in jobs))
        self.assertTrue(all(job["amap_company_url"].startswith("https://www.amap.com/search") for job in jobs))

    def test_smart_search_prioritizes_agent_when_python_and_agent_are_both_requested(self):
        jobs = build_smart_search_jobs(
            "AI Agent LLM应用 RAG Python FastAPI 数据处理 实习 初级",
            ["武汉"],
            3,
            candidate_profile="大三计科，想找 Python Agent 实习。",
        )

        self.assertGreaterEqual(len(jobs), 1)
        self.assertIn("Agent", jobs[0]["job_title"])

    def test_smart_search_keeps_python_internship_separate_from_agent(self):
        jobs = build_smart_search_jobs(
            "python实习",
            ["武汉"],
            3,
            candidate_profile="大三计科，学过 Python、Java、MySQL、Vue。",
        )

        self.assertGreaterEqual(len(jobs), 1)
        self.assertEqual(jobs[0]["job_title"], "Python实习生")
        self.assertFalse(any("Agent" in job["job_title"] or "大模型" in job["job_title"] for job in jobs))

    def test_deepseek_search_planner_parses_job_title_plan(self):
        response = Mock()
        response.raise_for_status.return_value = None
        response.json.return_value = {
            "choices": [
                {
                    "message": {
                        "content": json.dumps(
                            {
                                "search_intent": "全栈开发实习",
                                "recommended_titles": [
                                    {
                                        "job_title": "全栈开发实习生",
                                        "search_query": "全栈开发实习生 武汉",
                                        "reason": "简历包含 Vue、MySQL 和接口经验，适合前后端协作岗位。",
                                        "tags": ["全栈", "Vue", "MySQL", "实习"],
                                    }
                                ],
                            },
                            ensure_ascii=False,
                        )
                    }
                }
            ]
        }
        session = Mock()
        session.post.return_value = response
        client = DeepSeekClient(api_key="test-key", session=session)

        plan = client.plan_job_search(
            keywords="全栈",
            locations=["武汉"],
            job_type="internship",
            experience_level="entry-level",
            candidate_profile="大三计科，做过 Vue、MySQL、接口自动化和 Java 课程项目。",
            max_jobs=3,
        )

        self.assertEqual(plan[0]["job_title"], "全栈开发实习生")
        self.assertEqual(plan[0]["search_query"], "全栈开发实习生 武汉")
        self.assertIn("Vue", plan[0]["tags"])

    def test_smart_search_uses_ai_plan_for_fullstack_keyword(self):
        search_plan = [
            {
                "job_title": "全栈开发实习生",
                "search_query": "全栈开发实习生 武汉",
                "reason": "关键词是岗位名称，全栈应优先找全栈岗位。",
                "tags": ["全栈", "Vue", "MySQL", "实习"],
            },
            {
                "job_title": "Web全栈实习生",
                "search_query": "Web全栈实习生 武汉",
                "reason": "简历有 Web 和数据库基础。",
                "tags": ["Web", "全栈", "实习"],
            },
        ]

        jobs = build_smart_search_jobs(
            "全栈",
            ["武汉"],
            2,
            candidate_profile="大三计科，学过 Java、Python、Vue、MySQL。",
            job_type="internship",
            experience_level="entry-level",
            search_plan=search_plan,
        )

        self.assertEqual([job["job_title"] for job in jobs], ["全栈开发实习生", "Web全栈实习生"])
        self.assertTrue(all(job["source"] == "company_candidate" for job in jobs))
        self.assertTrue(all("city=101200100" in job["source_url"] for job in jobs))
        self.assertFalse(any(job["job_title"].startswith("Java") for job in jobs))

    def test_default_search_calls_ai_planner_before_generating_jobs(self):
        fake_deepseek = Mock()
        fake_deepseek.plan_job_search.return_value = [
            {
                "job_title": "全栈开发实习生",
                "search_query": "全栈开发实习生 武汉",
                "reason": "结合岗位名和简历筛选。",
                "tags": ["全栈", "实习"],
            }
        ]
        fake_deepseek.analyze_job.return_value = {
            "provider": "deepseek",
            "status": "ready",
            "match_score": 86,
            "summary": "岗位方向和简历匹配。",
        }
        fake_live_jobs = [
            {
                "source": "web_search",
                "source_url": "https://www.zhipin.com/web/geek/job?query=fullstack",
                "link_status": "needs_verification",
                "job_title": "全栈开发实习生",
                "company_name": "武汉真实科技有限公司",
                "job_location": "武汉",
                "salary": "150-200元/天",
                "job_description": "Vue MySQL 接口 实习",
                "tags": ["全栈", "实习"],
            }
        ]

        with patch("src.scraper.search.boss_scraper.search_web_real_company_jobs", return_value=fake_live_jobs):
            results = run_boss_deepseek_search(
                "全栈",
                ["武汉"],
                1,
                candidate_profile="大三计科，做过 Vue、MySQL 和接口项目。",
                job_type="internship",
                experience_level="entry-level",
                deepseek_client=fake_deepseek,
            )

        fake_deepseek.plan_job_search.assert_called_once()
        self.assertEqual(results[0]["job_title"], "全栈开发实习生")
        self.assertEqual(results[0]["ai_analysis"]["match_score"], 86)

    def test_default_search_uses_real_company_candidates_without_opening_boss(self):
        fake_deepseek = Mock()
        fake_deepseek.analyze_job.return_value = {
            "provider": "deepseek",
            "status": "ready",
            "match_score": 80,
            "summary": "岗位方向匹配",
        }
        fake_live_jobs = [
            {
                "source": "web_search",
                "source_url": "https://www.zhipin.com/web/geek/job?query=java",
                "link_status": "needs_verification",
                "job_title": "Java开发实习生",
                "company_name": "武汉真实软件有限公司",
                "job_location": "武汉",
                "salary": "150-200元/天",
                "job_description": "Java MySQL 实习",
                "tags": ["Java", "实习"],
            }
        ]

        with patch(
            "src.scraper.search.boss_scraper.search_boss_jobs",
            side_effect=AssertionError("BOSS browser should not open by default"),
        ), patch("src.scraper.search.boss_scraper.search_web_real_company_jobs", return_value=fake_live_jobs):
            results = run_boss_deepseek_search(
                "java",
                ["武汉"],
                2,
                candidate_profile="大三计科，学过 Java、MySQL、Vue 和接口自动化。",
                deepseek_client=fake_deepseek,
            )

        self.assertGreaterEqual(len(results), 1)
        self.assertEqual(results[0]["source"], "web_search")
        self.assertIn("Java", results[0]["job_title"])
        self.assertEqual(results[0]["company_name"], "武汉真实软件有限公司")
        self.assertTrue(results[0]["source_url"].startswith("http"))
        self.assertEqual(results[0]["ai_analysis"]["match_score"], 80)

    def test_python_search_does_not_expand_to_agent_when_keyword_is_python_internship(self):
        fake_deepseek = Mock()
        fake_deepseek.plan_job_search.return_value = [
            {
                "job_title": "Python开发实习生",
                "search_query": "Python开发实习生 武汉",
                "reason": "用户明确搜索 Python 实习，优先普通 Python 开发岗位。",
                "tags": ["Python", "实习"],
            }
        ]
        fake_deepseek.analyze_job.return_value = {
            "provider": "deepseek",
            "status": "ready",
            "match_score": 78,
            "summary": "Python 实习方向匹配",
        }
        fake_live_jobs = [
            {
                "source": "web_search",
                "source_url": f"https://www.zhipin.com/web/geek/job?query=python{i}",
                "link_status": "needs_verification",
                "job_title": title,
                "company_name": f"武汉真实科技{i}有限公司",
                "job_location": "武汉",
                "salary": "150-200元/天",
                "job_description": "Python 实习",
                "tags": ["Python", "实习"],
            }
            for i, title in enumerate(["Python开发实习生", "Python数据处理实习生", "Python爬虫实习生"], start=1)
        ]

        with patch("src.scraper.search.boss_scraper.search_web_real_company_jobs", return_value=fake_live_jobs):
            results = run_boss_deepseek_search(
                "python实习",
                ["武汉"],
                3,
                candidate_profile="大三计算机科学与技术，做过 Vue、MySQL、接口自动化和 Python 课程项目。",
                job_type="internship",
                experience_level="entry-level",
                deepseek_client=fake_deepseek,
            )

        self.assertEqual(len(results), 3)
        self.assertTrue(all("Agent" not in job["job_title"] for job in results))
        self.assertTrue(all(job["source"] == "web_search" for job in results))
        self.assertTrue(all("真实科技" in job["company_name"] for job in results))

    def test_deepseek_analysis_normalizes_required_sections_when_model_omits_them(self):
        session = Mock()
        response = Mock()
        response.raise_for_status.return_value = None
        response.json.return_value = {
            "choices": [{"message": {"content": '{"match_score": 70, "summary": "基础匹配", "recommendation": "建议投递"}'}}]
        }
        session.post.return_value = response

        client = DeepSeekClient(api_key="test-key", session=session)
        analysis = client.analyze_job(
            {
                "job_title": "Python开发实习生",
                "company_name": "可循智能",
                "job_description": "学过 Python 或 Java 基础，MySQL 数据库，学习过 HTML5。",
            },
            candidate_profile="大三计科，学过 Python、Java、MySQL。",
        )

        self.assertEqual(analysis["provider"], "deepseek")
        self.assertEqual(analysis["status"], "ready")
        self.assertIsInstance(analysis["skill_gaps"], list)
        self.assertIsInstance(analysis["resume_tips"], list)
        self.assertIsInstance(analysis["resume_rewrite_bullets"], list)
        self.assertIsInstance(analysis["interview_questions"], list)
        self.assertIn("Python开发实习生", analysis["self_introduction"])

    def test_extracts_resume_text_from_docx_upload(self):
        from docx import Document

        buffer = BytesIO()
        document = Document()
        document.add_paragraph("王同学")
        document.add_paragraph("项目经历：Python Agent 求职助手")
        document.save(buffer)

        client = TestClient(main_api.app)
        response = client.post(
            "/resume/extract",
            json={
                "file_name": "resume.docx",
                "content_base64": base64.b64encode(buffer.getvalue()).decode("ascii"),
            },
        )

        self.assertEqual(response.status_code, 200)
        text = response.json()["data"]["text"]
        self.assertIn("王同学", text)
        self.assertIn("Python Agent 求职助手", text)

    def test_blocked_boss_search_returns_actionable_result(self):
        fake_deepseek = Mock()
        fake_deepseek.analyze_job.return_value = {
            "provider": "deepseek",
            "status": "ready",
            "match_score": 76,
            "summary": "已切换到精选岗位",
        }

        with patch.dict(os.environ, {"BOSS_AUTOMATION_ENABLED": "true"}, clear=False), patch(
            "src.scraper.search.boss_scraper.search_boss_jobs",
            side_effect=RuntimeError("BOSS 连接被关闭"),
        ):
            results = run_boss_deepseek_search(
                "python实习",
                ["武汉"],
                1,
                deepseek_client=fake_deepseek,
            )

        self.assertEqual(results, [])

    def test_analyze_boss_text_endpoint_returns_structured_ai_result(self):
        client = TestClient(main_api.app)
        fake_analysis = {
            "provider": "deepseek",
            "match_score": 88,
            "summary": "岗位和 Python Agent 方向匹配。",
            "recommendation": "建议投递",
            "skill_gaps": ["FastAPI 项目深度", "LLM API 调用经验"],
            "resume_rewrite_bullets": [
                "基于 FastAPI 和 DeepSeek API 搭建求职 Agent，实现岗位匹配分析和简历优化建议。"
            ],
            "self_introduction": "面试官您好，我是计算机专业大三学生，做过 Python Agent 求职助手项目。",
            "interview_questions": ["你如何设计 Agent 工具调用流程？"],
        }
        fake_live_jobs = [
            {
                "source": "web_search",
                "source_url": "https://www.zhipin.com/web/geek/job?query=java",
                "link_status": "needs_verification",
                "job_title": "Java开发实习生",
                "company_name": "武汉真实软件有限公司",
                "job_location": "武汉",
                "salary": "150-200元/天",
                "job_description": "Java MySQL SpringBoot 实习",
                "tags": ["Java", "SpringBoot", "实习"],
            }
        ]

        with patch("main_api.DeepSeekClient") as client_class, patch(
            "main_api.search_web_real_company_jobs", return_value=fake_live_jobs
        ):
            client_class.return_value.analyze_job.return_value = fake_analysis
            response = client.post(
                "/boss/analyze-text",
                json={
                    "job_text": (
                        "职位：Python Agent 实习生\n"
                        "公司：武汉智能科技有限公司\n"
                        "地点：武汉 洪山区\n"
                        "薪资：150-200元/天\n"
                        "要求：Python、FastAPI、LLM API"
                    ),
                    "candidate_profile": "大三计科，熟悉 Python，正在补 Agent 项目。",
                },
            )

        self.assertEqual(response.status_code, 200)
        analyzed_job = client_class.return_value.analyze_job.call_args.args[0]
        candidate_profile = client_class.return_value.analyze_job.call_args.kwargs["candidate_profile"]
        data = response.json()["data"]
        self.assertEqual(analyzed_job["job_title"], "Python Agent 实习生")
        self.assertIn("正在补 Agent 项目", candidate_profile)
        self.assertEqual(data["job_title"], "Python Agent 实习生")
        self.assertEqual(data["company_name"], "武汉智能科技有限公司")
        self.assertEqual(data["job_location"], "武汉 洪山区")
        self.assertEqual(data["salary"], "150-200元/天")
        self.assertEqual(data["ai_analysis"]["match_score"], 88)
        self.assertIn("FastAPI 项目深度", data["ai_analysis"]["skill_gaps"])
        self.assertIn("求职 Agent", data["ai_analysis"]["resume_rewrite_bullets"][0])
        self.assertIn("面试官您好", data["ai_analysis"]["self_introduction"])

    def test_boss_search_endpoint_passes_candidate_profile_to_agent(self):
        temp_dir = tempfile.TemporaryDirectory()
        self.addCleanup(temp_dir.cleanup)
        client = TestClient(main_api.app)
        fake_jobs = [
            {
                "job_title": "Python Agent 实习生",
                "company_name": "武汉智能科技有限公司",
                "job_location": "武汉 洪山区",
                "source": "boss",
            }
        ]

        with patch("main_api.output_dir", temp_dir.name), patch(
            "main_api.search_history_file",
            os.path.join(temp_dir.name, "search_history.json"),
        ), patch("main_api.is_boss_automation_enabled", return_value=False), patch(
            "main_api.build_fast_boss_candidates", return_value=fake_jobs
        ) as search:
            response = client.post(
                "/search",
                json={
                    "keywords": "python实习",
                    "locations": ["武汉"],
                    "job_type": "internship",
                    "experience_level": "entry-level",
                    "max_jobs": 5,
                    "scrapers": ["boss"],
                    "candidate_profile": "大三计科，熟悉 FastAPI，想做 Agent。",
                },
                headers={"X-API-Key": main_api.API_KEY},
            )

        self.assertEqual(response.status_code, 200)
        self.assertIn("想做 Agent", search.call_args.kwargs["candidate_profile"])

    def test_fast_candidates_include_resume_analysis_and_match_selected_district(self):
        fake_live_jobs = [
            {
                "source": "web_search",
                "source_url": "https://www.zhipin.com/web/geek/job?query=python",
                "link_status": "needs_verification",
                "job_title": "Python开发实习生",
                "company_name": "武汉真实科技有限公司",
                "job_location": "武汉 汉阳区",
                "salary": "150-200元/天",
                "job_description": "Python FastAPI 实习",
                "tags": ["Python", "实习"],
            }
        ]
        with patch("main_api.search_web_real_company_jobs", return_value=fake_live_jobs):
            results = main_api.build_fast_boss_candidates(
                keywords="python agent",
                locations=["武汉 汉阳区"],
                max_jobs=3,
                candidate_profile="大三计科，学过 Python、Java、MySQL、Vue，想找开发实习。",
                job_type="internship",
                experience_level="entry-level",
            )

        self.assertGreater(len(results), 0)
        self.assertTrue(any("汉阳区" in job["job_location"] for job in results))
        self.assertTrue(all("ai_analysis" in job for job in results))
        self.assertTrue(all(job["ai_analysis"]["provider"] == "local-free" for job in results))
        self.assertTrue(all(job["ai_analysis"]["match_score"] is not None for job in results))

    def test_fast_candidates_do_not_fallback_to_generated_company_candidates(self):
        with patch("main_api.search_web_real_company_jobs", return_value=[]), patch(
            "main_api.build_legacy_candidate_jobs",
            side_effect=AssertionError("strict real search must not use generated company candidates"),
        ):
            results = main_api.build_fast_boss_candidates(
                keywords="python agent",
                locations=["武汉 东西湖区"],
                max_jobs=3,
                candidate_profile="大三计科，学过 Python、Java、MySQL、Vue，想找开发实习。",
                job_type="internship",
                experience_level="entry-level",
            )

        self.assertEqual(results, [])

    def test_strict_real_search_keeps_only_live_sources(self):
        fake_jobs = [
            {
                "source": "web_search",
                "source_url": "https://www.zhipin.com/web/geek/job?query=python",
                "link_status": "needs_verification",
                "job_title": "Python开发实习生",
                "company_name": "武汉真实科技有限公司",
                "job_location": "武汉 东西湖区",
                "salary": "150-200元/天",
                "job_description": "Python FastAPI 实习",
                "tags": ["Python", "实习"],
            }
        ]

        with patch("main_api.is_boss_automation_enabled", return_value=True), patch(
            "main_api.search_boss_jobs", return_value=[]
        ), patch("main_api.search_web_real_company_jobs", return_value=fake_jobs), patch(
            "main_api.search_observed_real_company_jobs",
            side_effect=AssertionError("strict real search must not use observed samples"),
        ), patch(
            "main_api.build_smart_search_jobs",
            side_effect=AssertionError("strict real search must not use generated candidates"),
        ):
            results = main_api.build_fast_boss_candidates(
                keywords="python",
                locations=["武汉 东西湖区"],
                max_jobs=3,
                candidate_profile="大三计科，学过 Python。",
                job_type="internship",
                experience_level="entry-level",
            )

        self.assertEqual(len(results), 1)
        self.assertEqual(results[0]["source"], "web_search")
        self.assertEqual(results[0]["company_name"], "武汉真实科技有限公司")

    def test_run_search_uses_real_company_web_results_instead_of_smart_placeholders(self):
        fake_web_jobs = [
            {
                "source": "web_search",
                "source_url": "https://www.zhipin.com/web/geek/job?query=fullstack",
                "link_status": "needs_verification",
                "job_title": "Full Stack Intern",
                "company_name": "Real Wuhan Tech",
                "job_location": "Wuhan",
                "salary": "platform",
                "job_description": "Python and Vue internship",
                "tags": ["Python", "Vue"],
            }
        ]
        deepseek = Mock()
        deepseek.plan_job_search.return_value = []
        deepseek.analyze_job.return_value = {"match_score": 80}

        with patch("src.scraper.search.boss_scraper.search_observed_real_company_jobs", return_value=[]), patch(
            "src.scraper.search.boss_scraper.search_web_real_company_jobs",
            return_value=fake_web_jobs,
        ):
            jobs = run_boss_deepseek_search(
                "fullstack",
                ["Wuhan"],
                3,
                candidate_profile="Python Vue MySQL",
                job_type="internship",
                experience_level="entry-level",
                deepseek_client=deepseek,
            )

        self.assertEqual(jobs[0]["company_name"], "Real Wuhan Tech")
        self.assertEqual(jobs[0]["source"], "web_search")
        self.assertNotEqual(jobs[0]["source"], "smart_search")

    def test_run_search_falls_back_to_smart_candidates_when_no_real_company_exists(self):
        deepseek = Mock()
        deepseek.plan_job_search.return_value = []
        deepseek.analyze_job.return_value = {"match_score": 60}

        with patch("src.scraper.search.boss_scraper.search_observed_real_company_jobs", return_value=[]), patch(
            "src.scraper.search.boss_scraper.search_web_real_company_jobs",
            return_value=[],
        ):
            jobs = run_boss_deepseek_search(
                "unknown role",
                ["Wuhan"],
                3,
                candidate_profile="Python Vue MySQL",
                job_type="internship",
                experience_level="entry-level",
                deepseek_client=deepseek,
        )

        self.assertEqual(jobs, [])

    def test_imported_job_keeps_frontend_and_backend_alias_fields(self):
        job = main_api.normalize_imported_job(
            {
                "title": "Python Intern",
                "company": "Real Company",
                "location": "Wuhan",
                "salary": "\ue000\ue001/Day",
                "link": "https://www.zhipin.com/job_detail/abc.html",
            },
            default_location="Wuhan",
        )

        self.assertEqual(job["job_title"], "Python Intern")
        self.assertEqual(job["title"], "Python Intern")
        self.assertEqual(job["name"], "Python Intern")
        self.assertEqual(job["company_name"], "Real Company")
        self.assertEqual(job["company"], "Real Company")
        self.assertEqual(job["job_location"], "Wuhan")
        self.assertEqual(job["location"], "Wuhan")
        self.assertNotIn("\ue000", job["salary"])

    def test_legacy_candidates_fill_requested_count_after_district_priority(self):
        observed_jobs = [
            {
                "job_title": "Python实习生",
                "company_name": "武汉云简科技",
                "job_location": "武汉 汉阳区",
                "salary": "150-200元/天",
                "job_description": "Python FastAPI 数据处理",
            },
            {
                "job_title": "Python开发实习生",
                "company_name": "武汉智简科技有限公司",
                "job_location": "武汉 江岸区",
                "salary": "150-200元/天",
                "job_description": "Python Web 开发",
            },
            {
                "job_title": "数据分析实习生",
                "company_name": "准星科技",
                "job_location": "武汉 江岸区",
                "salary": "100-150元/天",
                "job_description": "Python 数据分析",
            },
        ]
        smart_jobs = [
            {
                "job_title": "后端开发实习生",
                "company_name": "武汉简历云科技",
                "job_location": "武汉",
                "salary": "150-200元/天",
                "job_description": "Python 后端",
            },
            {
                "job_title": "Web开发实习生",
                "company_name": "武汉代码工坊",
                "job_location": "武汉",
                "salary": "120-180元/天",
                "job_description": "Python Web",
            },
        ]

        with patch("main_api.search_observed_real_company_jobs", return_value=observed_jobs), patch(
            "main_api.build_smart_search_jobs", return_value=smart_jobs
        ), patch("main_api.DeepSeekClient") as client_class:
            client_class.return_value.analyze_job.side_effect = [
                {"match_score": score, "summary": "测试匹配"}
                for score in [70, 68, 66, 64, 62]
            ]
            jobs = main_api.build_legacy_candidate_jobs(
                keywords="python实习",
                locations=["武汉 汉阳区"],
                max_jobs=5,
                candidate_profile="大三计科，学过 Python、Vue、MySQL。",
            )

        self.assertEqual(len(jobs), 5)
        self.assertEqual(jobs[0]["company_name"], "武汉云简科技")
        self.assertTrue(all(job["source"] == "company_candidate" for job in jobs))
        self.assertTrue(all(job.get("amap_company_url") for job in jobs))

    def test_fast_candidates_top_off_when_live_search_returns_too_few_jobs(self):
        live_jobs = [
            {
                "source": "web_search",
                "source_url": "https://www.zhipin.com/web/geek/job?query=python",
                "link_status": "needs_verification",
                "job_title": "Python FastAPI 实习生",
                "company_name": "武汉真实科技",
                "job_location": "武汉 汉阳区",
                "salary": "150-200元/天",
                "job_description": "Python FastAPI 数据处理",
            }
        ]
        observed_jobs = [
            {
                "job_title": "Python开发实习生",
                "company_name": "武汉云简科技",
                "job_location": "武汉 汉阳区",
                "salary": "150-200元/天",
                "job_description": "Python Web 开发",
            },
            {
                "job_title": "数据分析实习生",
                "company_name": "准星科技",
                "job_location": "武汉 江岸区",
                "salary": "100-150元/天",
                "job_description": "Python 数据分析",
            },
            {
                "job_title": "后端开发实习生",
                "company_name": "武汉简历云科技",
                "job_location": "武汉",
                "salary": "150-200元/天",
                "job_description": "Python 后端",
            },
            {
                "job_title": "Web开发实习生",
                "company_name": "武汉代码工坊",
                "job_location": "武汉",
                "salary": "120-180元/天",
                "job_description": "Python Web",
            },
        ]

        with patch("main_api.search_web_real_company_jobs", return_value=live_jobs), patch(
            "main_api.search_observed_real_company_jobs", return_value=observed_jobs
        ), patch("main_api.build_smart_search_jobs", return_value=[]), patch(
            "main_api.DeepSeekClient"
        ) as client_class:
            client_class.return_value.analyze_job.side_effect = [
                {"match_score": score, "summary": "测试匹配"}
                for score in [75, 72, 69, 66, 63]
            ]
            jobs = main_api.build_fast_boss_candidates(
                keywords="python实习",
                locations=["武汉 汉阳区"],
                max_jobs=5,
                candidate_profile="大三计科，学过 Python、Vue、MySQL。",
            )

        self.assertEqual(len(jobs), 1)
        self.assertTrue(any(job["company_name"] == "武汉真实科技" for job in jobs))
        self.assertEqual({job["source"] for job in jobs}, {"web_search"})

    def test_career_analyze_recommends_companies_and_learning_gaps_from_resume(self):
        client = TestClient(main_api.app)
        fake_analysis = {
            "summary": "简历更适合 Java 后端实习，已有 Java、MySQL、Vue 基础。",
            "best_fit_roles": ["Java开发实习生", "后端开发实习生"],
            "skill_gaps": ["SpringBoot 项目深度", "Redis 缓存实践"],
            "experience_gaps": ["缺少 3 个月以上实习经历"],
            "resume_fixes": ["把 Java/MySQL 项目放到简历前半部分"],
            "learning_plan": [
                {
                    "topic": "SpringBoot + MyBatis 项目",
                    "why": "多数 Java 实习任职要求会出现",
                    "platform_keywords": {
                        "bilibili": "SpringBoot MyBatis 实战 项目",
                        "baidu": "Java 实习 SpringBoot MyBatis 任职要求",
                        "douyin": "Java实习项目 SpringBoot",
                        "xiaohongshu": "Java实习 简历 项目"
                    },
                }
            ],
            "hot_requirements": ["Java 基础", "MySQL", "SpringBoot", "Vue"],
            "next_actions": ["先补一个 SpringBoot CRUD 项目"],
        }
        fake_live_jobs = [
            {
                "source": "web_search",
                "source_url": "https://www.zhipin.com/web/geek/job?query=java",
                "link_status": "needs_verification",
                "job_title": "Java开发实习生",
                "company_name": "武汉真实软件有限公司",
                "job_location": "武汉",
                "salary": "150-200元/天",
                "job_description": "Java MySQL SpringBoot 实习",
                "tags": ["Java", "SpringBoot", "实习"],
            }
        ]

        with patch("main_api.DeepSeekClient") as client_class, patch(
            "main_api.search_web_real_company_jobs", return_value=fake_live_jobs
        ):
            client_class.return_value.analyze_job.return_value = {
                "provider": "local-free",
                "match_score": 88,
                "summary": "测试匹配",
                "recommendation": "建议投递",
            }
            client_class.return_value.analyze_career_fit.return_value = fake_analysis
            response = client.post(
                "/career/analyze",
                json={
                    "candidate_profile": "大三计科，学过 Java、MySQL、Vue 和接口自动化。",
                    "target_role": "Java实习",
                    "location": "武汉",
                    "job_type": "internship",
                    "experience_level": "entry-level",
                    "requirement_text": "任职要求：有扎实 Java 基础，熟悉 MySQL，了解 SpringBoot。",
                    "max_recommendations": 3,
                },
            )

        self.assertEqual(response.status_code, 200)
        data = response.json()["data"]
        self.assertGreaterEqual(len(data["recommended_jobs"]), 1)
        self.assertEqual(data["recommended_jobs"][0]["company_name"], "武汉真实软件有限公司")
        self.assertEqual(data["career_analysis"]["skill_gaps"][0], "SpringBoot 项目深度")
        self.assertIn("SpringBoot", data["career_analysis"]["learning_plan"][0]["platform_keywords"]["bilibili"])
        client_class.return_value.analyze_career_fit.assert_called_once()

    def test_job_trends_aggregate_keyword_city_salary_and_company_type(self):
        jobs = [
            {
                "job_title": "Python开发实习生",
                "company_name": "武汉云简科技",
                "job_location": "武汉 洪山区",
                "salary": "150-200元/天",
                "job_description": "Python FastAPI 数据处理 MySQL",
            },
            {
                "job_title": "Python数据处理实习生",
                "company_name": "江城软件服务有限公司",
                "job_location": "武汉 江夏区",
                "salary": "100-150元/天",
                "job_description": "Python 爬虫 数据清洗 pandas",
            },
            {
                "job_title": "Java后端实习生",
                "company_name": "湖北汽车科技有限公司",
                "job_location": "武汉 洪山区",
                "salary": "150-200元/天",
                "job_description": "Java SpringBoot MySQL",
            },
        ]

        trends = main_api.build_job_trends(jobs, target_role="Python实习", location="武汉")

        self.assertEqual(trends["total_jobs"], 3)
        self.assertEqual(trends["keywords"][0]["name"], "Python")
        self.assertEqual(trends["cities"][0]["name"], "武汉 洪山区")
        self.assertEqual(trends["salary_ranges"][0]["name"], "150-200元/天")
        self.assertTrue(any(item["name"] == "软件服务" for item in trends["company_types"]))
        self.assertIn("市场最近更关注", trends["summary"])

    def test_career_analyze_response_includes_market_trends(self):
        client = TestClient(main_api.app)
        fake_live_jobs = [
            {
                "source": "web_search",
                "source_url": "https://example.com/python",
                "job_title": "Python开发实习生",
                "company_name": "武汉云简科技",
                "job_location": "武汉 洪山区",
                "salary": "150-200元/天",
                "job_description": "Python FastAPI 数据处理 MySQL",
            },
            {
                "source": "web_search",
                "source_url": "https://example.com/data",
                "job_title": "数据分析实习生",
                "company_name": "江城软件服务有限公司",
                "job_location": "武汉 江夏区",
                "salary": "100-150元/天",
                "job_description": "Python pandas SQL 数据分析",
            },
        ]

        with patch("main_api.build_fast_boss_candidates", return_value=fake_live_jobs), patch(
            "main_api.DeepSeekClient"
        ) as client_class:
            client_class.return_value.analyze_career_fit.return_value = {
                "summary": "优先匹配 Python开发实习生。",
                "best_fit_roles": ["Python开发实习生"],
                "skill_gaps": ["FastAPI 项目深度"],
                "experience_gaps": ["缺少完整 Python 项目"],
                "resume_fixes": ["突出 Python 项目"],
                "learning_plan": [],
                "hot_requirements": ["Python", "FastAPI"],
                "next_actions": ["补一个 FastAPI 项目"],
            }
            response = client.post(
                "/career/analyze",
                json={
                    "candidate_profile": "大三计科，学过 Python、MySQL、Vue。",
                    "target_role": "Python实习",
                    "location": "武汉",
                    "job_type": "internship",
                    "experience_level": "entry-level",
                    "max_recommendations": 5,
                },
            )

        self.assertEqual(response.status_code, 200)
        data = response.json()["data"]
        self.assertEqual(data["market_trends"]["total_jobs"], 2)
        self.assertEqual(data["market_trends"]["keywords"][0]["name"], "Python")
        self.assertIn("market_trends", data["career_analysis"])


    def test_imported_boss_jobs_create_visible_real_search_results(self):
        temp_dir = tempfile.TemporaryDirectory()
        self.addCleanup(temp_dir.cleanup)
        client = TestClient(main_api.app)

        with patch("main_api.output_dir", temp_dir.name), patch(
            "main_api.search_history_file",
            os.path.join(temp_dir.name, "search_history.json"),
        ), patch("main_api.DeepSeekClient") as client_class:
            client_class.return_value.analyze_job.return_value = {
                "provider": "local-free",
                "match_score": 82,
                "summary": "Imported real BOSS card matches the resume.",
            }
            response = client.post(
                "/imports/jobs",
                json={
                    "keywords": "python",
                    "locations": ["武汉 江夏区"],
                    "job_type": "internship",
                    "experience_level": "entry-level",
                    "max_jobs": 5,
                    "candidate_profile": "Python, Vue, MySQL student resume",
                    "jobs": [
                        {
                            "name": "Python开发实习生",
                            "company": "武汉云简科技",
                            "location": "武汉 江夏区",
                            "salary": "150-200元/天",
                            "link": "https://www.zhipin.com/job_detail/abc.html",
                            "tags": ["Python", "实习"],
                        }
                    ],
                },
            )

            self.assertEqual(response.status_code, 200)
            search_id = response.json()["search_id"]
            history_response = client.get("/search/history")
            results_response = client.get(f"/search/{search_id}")

        self.assertEqual(history_response.status_code, 200)
        self.assertEqual(history_response.json()["searches"][0]["job_count"], 1)
        self.assertEqual(results_response.status_code, 200)
        results = results_response.json()
        self.assertEqual(len(results), 1)
        self.assertEqual(results[0]["source"], "boss_import")
        self.assertEqual(results[0]["company_name"], "武汉云简科技")
        self.assertIn("amap.com/search", results[0]["amap_company_url"])
        self.assertEqual(results[0]["ai_analysis"]["match_score"], 82)

    def test_boss_collector_script_posts_page_jobs_to_import_endpoint(self):
        client = TestClient(main_api.app)

        response = client.get("/imports/boss-collector.js")

        self.assertEqual(response.status_code, 200)
        self.assertIn("javascript", response.headers["content-type"])
        script = response.text
        self.assertIn("/imports/jobs", script)
        self.assertIn("job-card-wrapper", script)
        self.assertIn("job_detail", script)

    def test_fast_candidates_prefer_recent_imported_real_jobs(self):
        temp_dir = tempfile.TemporaryDirectory()
        self.addCleanup(temp_dir.cleanup)

        imported = [
            {
                "source": "boss_import",
                "job_title": "Python intern",
                "company_name": "Wuhan Real Cloud Tech",
                "job_location": "Wuhan Jiangxia",
                "salary": "150-200/day",
                "job_description": "Python FastAPI data processing internship",
                "tags": ["Python", "FastAPI"],
            }
        ]
        Path(temp_dir.name, "job_import_20260703_120000.json").write_text(
            json.dumps(imported),
            encoding="utf-8",
        )

        with patch("main_api.output_dir", temp_dir.name), patch(
            "main_api.search_web_real_company_jobs", return_value=[]
        ), patch("main_api.is_boss_automation_enabled", return_value=False), patch(
            "main_api.DeepSeekClient"
        ) as client_class:
            client_class.return_value.analyze_job.return_value = {
                "provider": "local-free",
                "match_score": 88,
                "summary": "real imported job",
            }
            results = main_api.build_fast_boss_candidates(
                keywords="python",
                locations=["Wuhan"],
                max_jobs=5,
                candidate_profile="Python resume",
                job_type="internship",
                experience_level="entry-level",
            )

        self.assertEqual(len(results), 1)
        self.assertEqual(results[0]["source"], "boss_import")
        self.assertEqual(results[0]["company_name"], "Wuhan Real Cloud Tech")
        self.assertEqual(results[0]["ai_analysis"]["match_score"], 88)

    def test_imported_jobs_do_not_fill_with_unmatched_real_jobs(self):
        temp_dir = tempfile.TemporaryDirectory()
        self.addCleanup(temp_dir.cleanup)

        imported = [
            {
                "source": "boss_import",
                "job_title": "Python intern",
                "company_name": "Wuhan Python Tech",
                "job_location": "Wuhan",
                "salary": "150-200/day",
                "job_description": "Python FastAPI internship",
                "tags": ["Python"],
            },
            {
                "source": "boss_import",
                "job_title": "Social product intern",
                "company_name": "Wuhan Product Studio",
                "job_location": "Wuhan",
                "salary": "200/day",
                "job_description": "Community operations and product research",
                "tags": ["Product"],
            },
        ]
        Path(temp_dir.name, "job_import_20260703_130000.json").write_text(
            json.dumps(imported),
            encoding="utf-8",
        )

        with patch("main_api.output_dir", temp_dir.name), patch(
            "main_api.DeepSeekClient"
        ) as client_class:
            client_class.return_value.analyze_job.return_value = {
                "provider": "local-free",
                "match_score": 80,
                "summary": "matched",
            }
            results = main_api.load_recent_imported_jobs(
                keywords="python agent",
                locations=["Wuhan"],
                max_jobs=5,
                candidate_profile="Python resume",
            )

        self.assertEqual(len(results), 1)
        self.assertEqual(results[0]["company_name"], "Wuhan Python Tech")

    def test_imported_salary_icon_font_is_hidden(self):
        job = main_api.normalize_imported_job(
            {
                "name": "Python intern",
                "company": "Wuhan Real AI",
                "location": "Wuhan",
                "salary": "\ue036-\ue032\ue031K",
                "tags": ["Python"],
            },
            default_location="Wuhan",
        )

        self.assertEqual(job["salary"], "薪资到 BOSS 核验")

    def test_imported_company_name_is_repaired_from_tags(self):
        job = main_api.normalize_imported_job(
            {
                "name": "Python intern",
                "company": "武汉·江夏区·流芳",
                "location": "武汉 江夏区",
                "salary": "150-200/day",
                "tags": ["Wuhan Real AI", "武汉·江夏区·流芳"],
            },
            default_location="武汉",
        )

        self.assertEqual(job["company_name"], "Wuhan Real AI")
        self.assertEqual(job["job_location"], "武汉 江夏区")

    def test_career_analysis_uses_recent_imported_jobs_as_recommendations(self):
        temp_dir = tempfile.TemporaryDirectory()
        self.addCleanup(temp_dir.cleanup)
        client = TestClient(main_api.app)

        imported = [
            {
                "source": "boss_import",
                "job_title": "Python intern",
                "company_name": "Wuhan Career Tech",
                "job_location": "Wuhan",
                "salary": "150-200/day",
                "job_description": "Python MySQL internship",
                "tags": ["Python", "MySQL"],
            }
        ]
        Path(temp_dir.name, "job_import_20260703_121000.json").write_text(
            json.dumps(imported),
            encoding="utf-8",
        )

        with patch("main_api.output_dir", temp_dir.name), patch(
            "main_api.search_web_real_company_jobs", return_value=[]
        ), patch("main_api.is_boss_automation_enabled", return_value=False), patch(
            "main_api.DeepSeekClient"
        ) as client_class:
            client_class.return_value.analyze_job.return_value = {
                "provider": "local-free",
                "match_score": 79,
                "summary": "career imported job",
            }
            client_class.return_value.analyze_career_fit.return_value = {
                "summary": "Prefer Python internship",
                "best_fit_roles": ["Python intern"],
                "skill_gaps": ["FastAPI"],
                "experience_gaps": [],
                "resume_fixes": [],
                "learning_plan": [],
                "hot_requirements": ["Python"],
                "next_actions": [],
            }
            response = client.post(
                "/career/analyze",
                json={
                    "candidate_profile": "Python MySQL resume",
                    "target_role": "python",
                    "location": "Wuhan",
                    "job_type": "internship",
                    "experience_level": "entry-level",
                    "max_recommendations": 5,
                },
            )

        self.assertEqual(response.status_code, 200)
        data = response.json()["data"]
        self.assertEqual(data["recommended_jobs"][0]["source"], "boss_import")
        self.assertEqual(data["recommended_jobs"][0]["company_name"], "Wuhan Career Tech")


if __name__ == "__main__":
    unittest.main()
